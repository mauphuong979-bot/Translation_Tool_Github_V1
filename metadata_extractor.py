import re
from docx import Document
import translation_lib as tl
from collections import Counter

def format_vn_date_to_digit(vn_date_str):
    """
    Converts '31 tháng 12 năm 2025' to '31/12/2025'.
    Also handles single digits by padding with 0.
    """
    if not vn_date_str:
        return vn_date_str
    
    # Match: day tháng month năm year
    match = re.search(r"(\d{1,2})\s+th\u00e1ng\s+(\d{1,2})\s+n\u0103m\s+(\d{4})", vn_date_str)
    if match:
        day, month, year = match.groups()
        return f"{int(day):02d}/{int(month):02d}/{year}"
    return vn_date_str

def extract_metadata(file_stream):
    """
    Extracts metadata from a Word document stream.
    Returns a dictionary of found fields.
    """
    metadata = {
        "name_vn": None,
        "year_end": None,
        "report_date": None,
        "period_in": None,
        "period_in_2": None
    }
    
    try:
        doc = Document(file_stream)
        all_text = []
        # 1. Collect text from paragraphs
        for para in doc.paragraphs:
            text = tl.clean_text(para.text)
            if text:
                all_text.append(text)
        
        # 2. Collect text from tables (important for headers and metadata)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Some cells have multiple paragraphs, clean_text handles them
                    text = tl.clean_text(cell.text)
                    if text:
                        all_text.append(text)
        
        full_content = "\n".join(all_text)
        
        # 1. Company Name (Vietnamese)
        # Anchor phrase: (gọi tắt là “Công ty”) đệ trình
        # Regex handles: optional parentheses, various quote styles (“ ” " '), and flexible whitespace
        # We also ensure normalization to NFC matching the clean_text output.
        anchor_regex = re.compile(
            r"(?:\(?)\s*g\u1ecdi t\u1eaft l\u00e0\s+[\u201c\"']C\u00f4ng ty[\u201d\"']\s*(?:\)?)\s*\u0111\u1ec7 tr\u00ecnh",
            re.IGNORECASE
        )
        
        for para_text in all_text:
            match_anchor = anchor_regex.search(para_text)
            if match_anchor:
                # Extract the text before the anchor
                pre_text = para_text[:match_anchor.start()].strip()
                # Remove trailing parenthesis or spaces
                pre_text = re.sub(r'[\s(]+$', '', pre_text)
                
                # Find the last occurrence of "Công ty" in the pre_text to get the full name
                # regex: (Công ty ...)$
                name_match = re.search(r"(C\u00f4ng ty\s+.*)$", pre_text, re.IGNORECASE)
                if name_match:
                    metadata["name_vn"] = name_match.group(1).strip()
                    break
        
        # 2. Year-end date
        ye_pattern = re.compile(r"kết thúc ngày\s+(?:ngày\s+)?(\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})", re.IGNORECASE)
        ye_match = ye_pattern.search(full_content)
        if ye_match:
            metadata["year_end"] = format_vn_date_to_digit(ye_match.group(1).strip())
        else:
            # Fallback 1: Common in Auditor's opinion paragraphs
            ye_pattern_2 = re.compile(r"tại ngày\s+(\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})", re.IGNORECASE)
            ye_match_2 = ye_pattern_2.search(full_content)
            if ye_match_2:
                metadata["year_end"] = format_vn_date_to_digit(ye_match_2.group(1).strip())
            
        # 3. Reporting date
        rep_pattern_1 = re.compile(r"được lập ngày\s+(?:ngày\s+)?(\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})", re.IGNORECASE)
        rep_match_1 = rep_pattern_1.search(full_content)
        if rep_match_1:
            metadata["report_date"] = format_vn_date_to_digit(rep_match_1.group(1).strip())
        else:
            rep_pattern_2 = re.compile(r"Ngày\s+(\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})")
            all_dates = rep_pattern_2.findall(full_content)
            if all_dates:
                metadata["report_date"] = format_vn_date_to_digit(all_dates[-1].strip())


        # 5. Periods (in table) - Short format
        # Pattern: Từ xx/xx/20xx đến xx/xx/20xx
        short_period_pattern = re.compile(
            r"([Tt]\u1eeb\s+\d{1,2}/\d{1,2}/\d{4}\s+[\u0111\u00d0]\u1ebfn\s+\d{1,2}/\d{1,2}/\d{4})",
            re.IGNORECASE
        )
        all_short_periods = short_period_pattern.findall(full_content)
        if all_short_periods:
            # Sort by frequency to get Period 1 and Period 2
            counts = Counter(all_short_periods).most_common(2)
            metadata["period_in"] = counts[0][0].strip()
            if len(counts) > 1:
                metadata["period_in_2"] = counts[1][0].strip()

        # --- FINAL FALLBACK: Derive Year-end from Period if still missing ---
        if not metadata["year_end"]:
            
            # B. Try to extract the second date from period_in (Từ xx/xx/xxxx đến xx/xx/xxxx)
            if not metadata["year_end"] and metadata["period_in"]:
                s_dates = re.findall(r"(\d{1,2})/(\d{1,2})/(\d{4})", metadata["period_in"])
                if len(s_dates) >= 2:
                    d, m, y = s_dates[-1] # Take the last one (end date)
                    metadata["year_end"] = f"{int(d):02d}/{int(m):02d}/{y}"

    except Exception as e:
        print(f"Error extracting metadata: {e}")
        
    return metadata
