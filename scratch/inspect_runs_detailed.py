
import docx
from docx.oxml.ns import qn

def inspect_runs_detailed(docx_path):
    doc = docx.Document(docx_path)
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Only check cells that have digits
                if any(char.isdigit() for char in cell.text):
                    print(f"  Cell({r_idx}, {c_idx}): '{cell.text}'")
                    for p_idx, para in enumerate(cell.paragraphs):
                        print(f"    Para {p_idx}:")
                        for run_idx, run in enumerate(para.runs):
                            # Print run text and whether it's part of a field
                            is_field = False
                            # Check XML for field markers in this run
                            if run._element.xpath('.//w:fldChar | .//w:instrText'):
                                is_field = True
                            print(f"      Run {run_idx}: '{run.text}' [Field: {is_field}]")

if __name__ == "__main__":
    inspect_runs_detailed("number.docx")
