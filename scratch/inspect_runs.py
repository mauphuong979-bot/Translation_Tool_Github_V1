
import sys
import os
import io
from docx import Document

# Set terminal output to UTF-8
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

# Add the project directory to sys.path
BASE_DIR = r"d:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V3"
sys.path.append(BASE_DIR)

import metadata_extractor as mex
import translation_lib as tl

def inspect_runs():
    doc_path = os.path.join(BASE_DIR, "fs_Usha_2025_v1.docx")
    doc = Document(doc_path)
    
    signer1 = "Phạm Phú Quí"
    signer2 = "Đỗ Thị Ngân Trâm"
    
    print(f"Inspecting runs for '{signer1}' and '{signer2}'...")
    
    def check_container(container, name):
        for para in container.paragraphs:
            if name in para.text:
                print(f"\nFound '{name}' in paragraph: '{para.text}'")
                print("Runs:")
                for i, run in enumerate(para.runs):
                    print(f"  Run {i}: '{run.text}'")
                    
    # Check body
    check_container(doc, signer1)
    check_container(doc, signer2)
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                check_container(cell, signer1)
                check_container(cell, signer2)

if __name__ == "__main__":
    inspect_runs()
