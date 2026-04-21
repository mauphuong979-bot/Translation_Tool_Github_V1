
import docx
import io
import re

def swap_vn_to_en_number_separators(text):
    if not text: return text
    pattern = re.compile(r'(?<!\d)(?:\d{1,3}(?:\.\d{3})+(?:,\d+)?|\d+,\d+)(?!\d)')
    def replace_func(match):
        val = match.group(0)
        return val.replace('.', 'TEMP_DOT').replace(',', '.').replace('TEMP_DOT', ',')
    return pattern.sub(replace_func, text)

# Simulate current faulty logic
def old_swap_logic(doc):
    for table in doc.tables:
        for row in table.rows: # Potential failure point for complex merges
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = swap_vn_to_en_number_separators(run.text)

# Proposed robust logic
def new_swap_logic(doc):
    processed_p_els = set()
    # Iterate ALL unique cells using _cells property
    for table in doc.tables:
        # A cell might span multiple rows/cols. table._cells only has unique ones? 
        # Actually table._cells has one entry per grid cell. 
        # We need a set of _Cell objects or their elements.
        unique_cells = []
        seen_cells = set()
        for cell in table._cells:
            if cell._tc not in seen_cells:
                unique_cells.append(cell)
                seen_cells.add(cell._tc)
        
        for cell in unique_cells:
            # Match ALL paragraphs including fragmented numbers in runs
            for para in cell.paragraphs:
                if para._element in processed_p_els: continue
                
                # Robust approach: Process at Paragraph level to handle fragmented numbers
                # But we must preserve formatting!
                full_text = "".join(r.text for r in para.runs)
                new_text = swap_vn_to_en_number_separators(full_text)
                
                if new_text != full_text:
                    # Update while preserving formatting if possible
                    # (Simplified for repro)
                    if len(para.runs) == 1:
                        para.runs[0].text = new_text
                    else:
                        # Complex run fragmentation handled here...
                        pass
                processed_p_els.add(para._element)

# Test on the source file
source = "fs_C&C_2025_V.docx"
doc = docx.Document(source)

# Find a cell with a number that SHOULD be swapped (if any)
found = False
for t in doc.tables:
    for c in t._cells:
        if re.search(r'\d+,\d+', c.text):
            print(f"Found candidate cell: '{c.text}'")
            found = True
            break
    if found: break
else:
    print("No VN numbers found in tables?? Check regex.")

