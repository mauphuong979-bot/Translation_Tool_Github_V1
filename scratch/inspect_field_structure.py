
import docx
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def inspect_para_xml(para):
    print(f"Para text: '{para.text}'")
    print("Runs XML structure:")
    for i, r in enumerate(para.runs):
        xml = r._element.xml
        # Just print if it has field markers
        has_field = "fldChar" in xml or "instrText" in xml or "fldSimple" in xml
        print(f"  Run[{i}]: text='{r.text}' has_field={has_field}")

doc = docx.Document("fs_C&C_2025_V.docx")
print("Searching for paragraphs with fields and numbers...")
for p in doc.paragraphs:
    # Check if it has a number and likely a field
    if any(c.isdigit() for c in p.text):
        # Use our has_fields logic
        element = p._element
        fields = element.xpath('.//*[local-name()="fldSimple" or local-name()="fldChar" or local-name()="instrText"]')
        if fields:
            inspect_para_xml(p)
            break

for t in doc.tables:
    for cell in t._cells:
        if any(c.isdigit() for c in cell.text):
            for p in cell.paragraphs:
                element = p._element
                fields = element.xpath('.//*[local-name()="fldSimple" or local-name()="fldChar" or local-name()="instrText"]')
                if fields:
                    print("\n--- Found field in table cell ---")
                    inspect_para_xml(p)
                    # Show XML for the first run that has the number
                    for r in p.runs:
                        if any(c.isdigit() for c in r.text):
                            print(f"Full XML of a number run: {r._element.xml}")
                            break
                    sys.exit(0)
