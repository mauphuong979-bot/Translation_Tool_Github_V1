
import docx
import sys
import io
import os
import unicodedata

# Set stdout to utf-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Import the library
import translation_lib as tl

def test_fragmentation():
    doc = docx.Document()
    p = doc.add_paragraph()
    # Create fragmented number: "1.234", "567", ",89"
    p.add_run("Số tiền là: ")
    p.add_run("1.234").bold = True
    p.add_run(".")
    p.add_run("567")
    p.add_run(",89").italic = True
    
    print(f"Original text: '{p.text}'")
    print(f"Runs: {[r.text for r in p.runs]}")
    
    # Run swap
    tl.apply_financial_number_formatting(doc, "E")
    
    # Swapped result should be "1,234,567.89"
    print(f"Swapped text: '{p.text}'")
    print(f"Swapped Runs: {[r.text for r in p.runs]}")
    # Check if formatting is preserved (at least in one run)
    for i, r in enumerate(p.runs):
        print(f"  Run[{i}]: '{r.text}' Bold={r.bold} Italic={r.italic}")

test_fragmentation()
