
import docx
import sys
import io
import os

# Set stdout to utf-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Import the library
import translation_lib as tl

def test_field_preservation():
    doc = docx.Document()
    p = doc.add_paragraph()
    
    # Create simulated field structure: 
    # [BeginChar] [instrText: LINK...] [SeparateChar] [Text: 1.234,56] [EndChar]
    p.add_run("Số dư: ")
    # Manually add field XML elements to simulate a link
    r_begin = p.add_run()
    r_begin._element.append(docx.oxml.shared.OxmlElement('w:fldChar'))
    r_begin._element.xpath('.//w:fldChar')[0].set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
    
    r_instr = p.add_run(" LINK Excel.Sheet.12 ... ")
    # Wrap instr text in instrText element
    instr_el = docx.oxml.shared.OxmlElement('w:instrText')
    instr_el.text = r_instr.text
    r_instr._element.remove(r_instr._element.xpath('.//w:t')[0])
    r_instr._element.append(instr_el)
    
    r_sep = p.add_run()
    r_sep._element.append(docx.oxml.shared.OxmlElement('w:fldChar'))
    r_sep._element.xpath('.//w:fldChar')[0].set(docx.oxml.ns.qn('w:fldCharType'), 'separate')
    
    r_text = p.add_run("1.234,56")
    
    r_end = p.add_run()
    r_end._element.append(docx.oxml.shared.OxmlElement('w:fldChar'))
    r_end._element.xpath('.//w:fldChar')[0].set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    
    print(f"Original text: '{p.text}'")
    print(f"Has Fields: {tl.has_fields(p)}")
    
    # Run swap
    tl.apply_financial_number_formatting(doc, "E")
    
    print(f"Swapped text: '{p.text}'")
    
    # Verify the field elements still exist
    xml = p._element.xml
    markers = ["fldChar", "instrText", "separate", "begin", "end"]
    all_present = all(m in xml for m in markers)
    print(f"Field markers preserved: {all_present}")
    if not all_present:
        print("MISSING MARKERS in XML!")
        # print(xml)

test_field_preservation()
