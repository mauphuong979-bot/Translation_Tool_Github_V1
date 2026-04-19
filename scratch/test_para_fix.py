import os
import sys
from docx import Document
import pandas as pd
import json

# Add current dir to path to import translation_lib
sys.path.append(os.path.abspath('.'))
import translation_lib as tl

def test_para_fix():
    # 1. Setup metadata
    metadata = {
        "name_vn": "Công ty TNHH Ito En Việt Nam",
        "name_trans": "Ito En Vietnam Co., Ltd",
        "year_end": "31/12/2025",
        "report_date": "20/03/2026",
    }
    
    # 2. Setup Para Map
    para_map = tl.load_para_template_map()
    
    # 3. Test Cases (Sample paragraphs from the user's report)
    test_cases = [
        "Tổng Giám đốc Công ty TNHH Ito En Việt Nam (gọi tắt là “Công ty”) đệ trình báo cáo này cùng với báo cáo tài chính đã được kiểm toán cho năm tài chính kết thúc ngày 31 tháng 12 năm 2025.",
        "Tôi, Murakami Hiroaki - Tổng Giám đốc, tuyên bố rằng báo cáo tài chính kèm theo đã được soạn lập đúng đắn, phù hợp với các Chuẩn mực kế toán Việt Nam và Chế độ kế toán doanh nghiệp Việt Nam và phản ánh trung thực và hợp lý tình hình tài chính của Công ty tại thời điểm ngày 31 tháng 12 năm 2025 và kết quả hoạt động kinh doanh và lưu chuyển tiền tệ cho năm tài chính kết thúc cùng ngày.",
        "Chúng tôi đã kiểm toán báo cáo tài chính kèm theo của Công ty TNHH Ito En Việt Nam (gọi tắt là “Công ty”), được lập ngày 20 tháng 03 năm 2026, từ trang 3 đến trang 18, bao gồm Bảng cân đối kế toán tại ngày 31 tháng 12 năm 2025, Báo cáo kết quả hoạt động kinh doanh, Báo cáo lưu chuyển tiền tệ cho năm tài chính kết thúc cùng ngày và Bản thuyết minh báo cáo tài chính."
    ]
    
    doc = Document()
    for text in test_cases:
        doc.add_paragraph(text)

    # 4. RUN PROCESS (Simplified steps)
    target_col = "E"
    
    print("\n--- BEFORE PROCESSING --- (Omitted VN output)")

    # STEP 3: ParaTemplate
    tl.apply_paragraph_templates(doc, para_map, target_col)
    
    # STEP 4: Dictionary (to fill metadata tags)
    v3_df = tl.load_and_fill_v3_dictionary(metadata)
    translation_map = dict(zip(v3_df['Vietnamese'], v3_df[target_col]))
    prepared_list = tl.prepare_translation_list(translation_map)
    tl._process_container(doc, prepared_list)
    
    # STEP 10: Highlight
    tl.highlight_vietnamese_text(doc)
    
    print("\n--- AFTER PROCESSING ---")
    for i, p in enumerate(doc.paragraphs):
        # Scan for highlights
        highlights = [run.font.highlight_color for run in p.runs if run.font.highlight_color]
        if highlights:
            print(f"P{i+1}: [!] Has highlights: {len(highlights)} runs")
        
        # Check if contains placeholders (yellow highlighted expected)
        if "[" in p.text and "]" in p.text:
            # Safely print part of the text if it's ascii-ish or just report finding
            print(f"P{i+1}: Contains likely placeholders (found '[' and ']')")
        
        # Check specific words known to be translated/replaced
        if "members of the Board" in p.text or "We have audited" in p.text or "I, [" in p.text:
            print(f"P{i+1}: CORRECT - Found expected Template English fragments.")
        else:
            print(f"P{i+1}: WARNING - Template fragments NOT found.")

    # Save for manual inspection
    os.makedirs("scratch", exist_ok=True)
    doc.save("scratch/test_para_fix_result.docx")
    print("\nResult saved to scratch/test_para_fix_result.docx")

if __name__ == "__main__":
    test_para_fix()
