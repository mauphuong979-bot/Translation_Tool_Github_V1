
import sys
import os
from docx import Document
import re

# Add root to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import translation_lib as tl

def main():
    doc_path = r'd:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V2\test_translated.docx'
    
    if not os.path.exists(doc_path):
        return

    doc = Document(doc_path)
    remaining_vn = []
    
    def check_text(text, context):
        if tl.contains_vietnamese(text):
            remaining_vn.append({"text": text, "context": context})

    for i, para in enumerate(doc.paragraphs):
        check_text(para.text, f"Para {i+1}")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                check_text(cell.text, f"T{t_idx+1} C({r_idx+1},{c_idx+1})")

    with open('scratch/verify_results.txt', 'w', encoding='utf-8') as f:
        f.write(f"Found {len(remaining_vn)} items still containing Vietnamese.\n\n")
        for item in remaining_vn:
            f.write(f"[{item['context']}] '{item['text']}'\n")

if __name__ == "__main__":
    main()
