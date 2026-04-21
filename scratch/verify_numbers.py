
import docx
import re

def count_patterns(doc):
    vn_num_pattern = re.compile(r'\d+(?:\.\d{3})*,\d+')
    en_num_pattern = re.compile(r'\d+(?:,\d{3})*\.\d+')
    
    vn_count = 0
    en_count = 0
    for t in doc.tables:
        for c in t._cells:
            if vn_num_pattern.search(c.text): vn_count += 1
            if en_num_pattern.search(c.text): en_count += 1
    return vn_count, en_count

doc_orig = docx.Document("fs_C&C_2025_V.docx")
vn_orig, en_orig = count_patterns(doc_orig)

doc_res = docx.Document("scratch/verified_output.docx")
vn_res, en_res = count_patterns(doc_res)

print(f"ORIGINAL: VN={vn_orig}, EN={en_orig}")
print(f"RESULT:   VN={vn_res}, EN={en_res}")
