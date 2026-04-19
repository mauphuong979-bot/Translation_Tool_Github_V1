
import sys
import os
import pandas as pd
from docx import Document
import unicodedata
import re

# Add root to path to import translation_lib
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import translation_lib

def main():
    doc_path = r'd:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V2\test.docx'
    dict_path = r'd:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V2\Dictionary_v3.xlsx'
    
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return

    print(f"--- Analyzing {os.path.basename(doc_path)} ---")
    
    # 1. Load Dictionary
    df_dict = pd.read_excel(dict_path).astype(str)
    # Build a lookup map of cleaned VN -> E
    dict_map = {}
    for _, row in df_dict.iterrows():
        vn = translation_lib.clean_text(row['Vietnamese'])
        en = translation_lib.clean_text(row['E'])
        if vn:
            dict_map[vn] = en
            
    print(f"Loaded {len(dict_map)} dictionary entries.")

    # 2. Load Document
    doc = Document(doc_path)
    
    # 3. Process and find issues
    # We will look for Vietnamese text that is NOT translated
    
    untranslated = []
    
    def log_untranslated(text, context):
        cleaned = translation_lib.clean_text(text)
        if not cleaned: return
        
        # Check if it has Vietnamese characters
        if translation_lib.contains_vietnamese(cleaned):
            # Check if it's in the dictionary
            if cleaned in dict_map:
                untranslated.append({
                    'text': text,
                    'cleaned': cleaned,
                    'context': context,
                    'reason': 'Found in dictionary but NOT replaced',
                    'dict_val': dict_map[cleaned]
                })
            else:
                # Try partial matches or case-insensitive
                found_partial = False
                for vn_key in dict_map:
                    if vn_key.lower() == cleaned.lower():
                        untranslated.append({
                            'text': text,
                            'cleaned': cleaned,
                            'context': context,
                            'reason': 'Case mismatch (Dictionary is different case)',
                            'dict_key': vn_key,
                            'dict_val': dict_map[vn_key]
                        })
                        found_partial = True
                        break
                
                if not found_partial:
                    untranslated.append({
                        'text': text,
                        'cleaned': cleaned,
                        'context': context,
                        'reason': 'Not found in dictionary',
                        'dict_val': 'N/A'
                    })

    # Scan Body
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            log_untranslated(para.text, f"Paragraph {i+1}")
            
    # Scan Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    log_untranslated(cell.text, f"Table {t_idx+1} Cell({r_idx+1}, {c_idx+1})")

    # Write results to a file (UTF-8)
    with open('scratch/debug_results.txt', 'w', encoding='utf-8') as f:
        f.write(f"Found {len(untranslated)} Vietnamese segments.\n\n")
        for item in untranslated:
            f.write(f"Context: {item['context']}\n")
            f.write(f"Text: '{item['text']}'\n")
            f.write(f"Cleaned: '{item['cleaned']}'\n")
            f.write(f"Reason: {item['reason']}\n")
            if 'dict_key' in item:
                f.write(f"Dict Key: '{item['dict_key']}'\n")
            f.write(f"Dict Val: '{item['dict_val']}'\n")
            f.write("-" * 40 + "\n")

    print(f"Results written to scratch/debug_results.txt. Found {len(untranslated)} items.")

if __name__ == "__main__":
    main()
