import os
import json
from docx import Document
import sys

# Set encoding for safe printing
sys.stdout.reconfigure(encoding='utf-8')

def extract_content(file_path):
    doc = Document(file_path)
    content = []
    
    content.append("--- MAIN BODY ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            content.append(f"P{i}: {p.text}")
    
    content.append("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        content.append(f"Table {i}:")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    # Avoid repeating same cell if merged
                    content.append(f"  T{i}R{r_idx}C{c_idx}: {cell.text.replace('\n', ' ')}")
                    
    with open('debug_doc_content.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(content))

if __name__ == "__main__":
    extract_content('fs_Itoen_2025_v - r2.docx')
    print("Extracted content to debug_doc_content.txt")
