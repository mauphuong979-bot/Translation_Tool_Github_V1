
import docx
import re
doc = docx.Document("fs_C&C_2025_V.docx")
en_num_pattern = re.compile(r'\d+(?:,\d{3})*\.\d+')
vn_num_pattern = re.compile(r'\d+(?:\.\d{3})*,\d+')

print("Samples from ORIGINAL (found as 'EN style' by regex):")
count = 0
for t in doc.tables:
    for c in t._cells:
        m = en_num_pattern.search(c.text)
        if m:
            print(f"Cell: '{c.text}' -> Match: '{m.group(0)}'")
            count += 1
            if count > 10: break
    if count > 10: break

print("\nSamples from ORIGINAL (found as 'VN style' by regex):")
count = 0
for t in doc.tables:
    for c in t._cells:
        m = vn_num_pattern.search(c.text)
        if m:
            print(f"Cell: '{c.text}' -> Match: '{m.group(0)}'")
            count += 1
            if count > 10: break
    if count > 10: break
