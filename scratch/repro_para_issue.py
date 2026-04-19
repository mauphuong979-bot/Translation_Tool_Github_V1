import os
import sys
from docx import Document
import pandas as pd
import json

# Add current dir to path to import translation_lib
sys.path.append(os.path.abspath('.'))
import translation_lib as tl

def repro():
    # 1. Load data
    para_map = tl.load_para_template_map()
    print(f"Loaded {len(para_map)} para templates.")
    
    # 2. Simulate processing of P35, P60, P80
    test_cases = [
        "Tổng Giám đốc Công ty TNHH Ito En Việt Nam (gọi tắt là “Công ty”) đệ trình báo cáo này cùng với báo cáo tài chính đã được kiểm toán cho năm tài chính kết thúc ngày 31 tháng 12 năm 2025.",
        "Tôi, Murakami Hiroaki - Tổng Giám đốc, tuyên bố rằng báo cáo tài chính kèm theo đã được soạn lập đúng đắn, phù hợp với các Chuẩn mực kế toán Việt Nam và Chế độ kế toán doanh nghiệp Việt Nam và phản ánh trung thực và hợp lý tình hình tài chính của Công ty tại thời điểm ngày 31 tháng 12 năm 2025 và kết quả hoạt động kinh doanh và lưu chuyển tiền tệ cho năm tài chính kết thúc cùng ngày.",
        "Chúng tôi đã kiểm toán báo cáo tài chính kèm theo của Công ty TNHH Ito En Việt Nam (gọi tắt là “Công ty”), được lập ngày 20 tháng 03 năm 2026, từ trang 3 đến trang 18, bao gồm Bảng cân đối kế toán tại ngày 31 tháng 12 năm 2025, Báo cáo kết quả hoạt động kinh doanh, Báo cáo lưu chuyển tiền tệ cho năm tài chính kết thúc cùng ngày và Bản thuyết minh báo cáo tài chính."
    ]
    
    # Create a dummy doc for testing
    doc = Document()
    for text in test_cases:
        doc.add_paragraph(text)
        
    print("\n--- BEFORE PARA TEMPLATE ---")
    for p in doc.paragraphs:
        print(f"[{p.text}]")
        
    # Apply templates
    tl.apply_paragraph_templates(doc, para_map, "E")
    
    print("\n--- AFTER PARA TEMPLATE (Current Logic) ---")
    for p in doc.paragraphs:
        print(f"[{p.text}]")

if __name__ == "__main__":
    repro()
