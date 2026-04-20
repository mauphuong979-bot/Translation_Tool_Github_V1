import sys
import os
from docx import Document
import re

# Add project root to path to import translation_lib
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import translation_lib

def test_abbreviation():
    doc = Document()
    
    # 1. Add a table with various month formats
    table = doc.add_table(rows=5, cols=1)
    table.cell(0, 0).text = "January 2023"        # Title Case -> Jan
    table.cell(1, 0).text = "FEBRUARY 2023"       # ALL CAPS -> FEBRUARY
    table.cell(2, 0).text = "march and april"    # lowercase -> mar and apr
    table.cell(3, 0).text = "September 15, 2024" # Title Case -> Sep
    table.cell(4, 0).text = "01/12/2023"         # Date format -> handled by format_dates_in_tables
    
    # 2. Add a paragraph outside stable (should NOT be abbreviated)
    doc.add_paragraph("December is the last month.")
    
    print("--- Before Processing ---")
    for i, row in enumerate(table.rows):
        p = row.cells[0].paragraphs[0]
        print(f"Cell {i}: '{p.text}' (Runs: {len(p.runs)})")
    print(f"Para: {doc.paragraphs[0].text}")
    
    # 3. Process
    # First apply date formatting (Step 5)
    translation_lib.format_dates_in_tables(doc, target_col="E")
    # Then apply English month abbreviation (Step 5.5)
    translation_lib.abbreviate_english_months_in_tables(doc)
    
    print("\n--- After Processing ---")
    results = []
    for i, row in enumerate(table.rows):
        txt = row.cells[0].text
        print(f"Cell {i}: {txt}")
        results.append(txt)
    
    para_txt = doc.paragraphs[0].text
    print(f"Para: {para_txt}")
    
    # 4. Assertions
    assert "Jan 2023" in results[0]
    assert "FEBRUARY 2023" in results[1] # Should stay as is
    assert "mar and apr" in results[2]
    assert "Sep 15, 2024" in results[3]
    assert "Dec 01, 2023" in results[4] # From format_dates_in_tables
    assert "December is the last month." == para_txt # Outside table, should stay full
    
    print("\nTest Passed!")

if __name__ == "__main__":
    try:
        test_abbreviation()
    except Exception as e:
        print(f"\nTest Failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
