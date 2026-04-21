
import docx
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    doc = docx.Document("fs_C&C_2025_V.docx")
    print("Searching for Form Indicator text in TABLES...")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text.upper()
                if ("MẪU" in t or "MÃ" in t or "FORM" in t) and "DN" in t:
                    print(f"Table[{ti}] Row[{ri}] Cell[{ci}]: '{cell.text}'")
                    for pi, p in enumerate(cell.paragraphs):
                        print(f"  Para[{pi}]: '{p.text}'")
                        for run in p.runs:
                            print(f"    Run: '{run.text}'")
except Exception as e:
    print(f"Error: {e}")
