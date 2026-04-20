
import sys
import os
import io
import pandas as pd
from docx import Document

# Set terminal output to UTF-8
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

# Add the project directory to sys.path
BASE_DIR = r"d:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V3"
sys.path.append(BASE_DIR)

import metadata_extractor as mex
import translation_lib as tl
from processor import process_financial_report

def run_test():
    doc_path = os.path.join(BASE_DIR, "fs_Usha_2025_v1.docx")
    print(f"Testing with: {doc_path}")
    
    with open(doc_path, "rb") as f:
        file_stream = io.BytesIO(f.read())
    
    # 1. Extract Metadata
    metadata = mex.extract_metadata(file_stream)
    print("\nExtracted Metadata:")
    for k, v in metadata.items():
        print(f"  {k}: {v}")
    
    # 2. Prepare process settings (mimic UI)
    process_settings = {
        "unicode": True,
        "clean_v": True,
        "para_template": True,
        "dictionary": True,
        "dual_font": True,
        "number_swap": True,
        "table_size": True,
        "date_format": True,
        "textbox": True,
        "signer_accents": True,
        "highlight": True,
        "suggestion": True
    }
    
    # 3. Load and fill dictionary (mimic app.py)
    metadata_for_tags = {
        "name_vn": metadata.get("name_vn") or "",
        "name_trans": "Usha (Thailand) Limited", 
        "year_end": metadata.get("year_end") or "",
        "report_date": metadata.get("report_date") or "",
        "period_in": metadata.get("period_in") or "",
        "period_in_2": metadata.get("period_in_2") or "",
        "signer_1": metadata.get("signer_1") or "",
        "signer_2": metadata.get("signer_2") or "",
        "signer_3": metadata.get("signer_3") or ""
    }
    
    v3_df = tl.load_and_fill_v3_dictionary(metadata_for_tags)
    target_col = "E"
    translation_map = dict(zip(v3_df['Vietnamese'], v3_df[target_col]))
    
    processor_metadata = metadata_for_tags

    # 4. Run Process
    file_stream.seek(0)
    processed_file, msg = process_financial_report(
        file_stream,
        metadata=processor_metadata,
        translation_map=translation_map,
        case_threshold=30,
        target_col=target_col,
        process_settings=process_settings
    )
    
    print(f"\nProcessing Result: {msg}")
    
    if processed_file:
        output_path = os.path.join(BASE_DIR, "scratch", "test_output_Usha.docx")
        if not os.path.exists(os.path.dirname(output_path)):
            os.makedirs(os.path.dirname(output_path))
        with open(output_path, "wb") as f:
            f.write(processed_file.getvalue())
        print(f"Saved output to: {output_path}")
        
        # Verify signer replacement in the output doc
        doc = Document(output_path)
        all_text = []
        for p in doc.paragraphs: all_text.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    all_text.append(c.text)
        
        full_content = "\n".join(all_text)
        
        print("\nVerifying Signers in output:")
        for i in range(1, 4):
            signer = metadata.get(f"signer_{i}")
            if signer:
                print(f"  Signer {i} (Accented): '{signer}'")
                if signer in full_content:
                    print(f"    STILL PRESENT (Accented)!")
                else:
                    print(f"    Replaced/Not found.")
                
                unaccented = tl.remove_accents(signer)
                if unaccented in full_content:
                    print(f"    FOUND (Unaccented): '{unaccented}'")
                else:
                    print(f"    NOT FOUND (Unaccented): '{unaccented}'")

if __name__ == "__main__":
    run_test()
