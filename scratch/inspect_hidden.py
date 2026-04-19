
import sys
import os
from docx import Document
import unicodedata

# Reconfigure stdout to use UTF-8
if sys.stdout.encoding != 'utf-8':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'replace')
    except:
        pass

doc_path = r'd:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V2\test.docx'
doc = Document(doc_path)

def inspect_for_hidden_chars(text, label, f):
    f.write(f"--- {label} ---\n")
    f.write(f"Original: {repr(text)}\n")
    for char in text:
        if ord(char) > 127 or ord(char) < 32:
            try:
                name = unicodedata.name(char, 'UNKNOWN')
            except:
                name = 'UNKNOWN'
            f.write(f"  Char: {repr(char)} | Code: {ord(char)} | Name: {name}\n")
    f.write("\n")

with open('scratch/hidden_chars.txt', 'w', encoding='utf-8') as f:
    # Inspect 'CHỈ TIÊU' in T1 C1,1
    try:
        inspect_for_hidden_chars(doc.tables[0].cell(0, 0).text, "T1 C1,1", f)
    except: pass
    
    # Inspect '6.	Doanh thu hoạt động tài chính' in T1 C13,1
    try:
        inspect_for_hidden_chars(doc.tables[0].cell(12, 0).text, "T1 C13,1", f)
    except: pass

    # Inspect 'Tổng Giám đốc' in T2 C2,1
    try:
        inspect_for_hidden_chars(doc.tables[1].cell(1, 0).text, "T2 C2,1", f)
    except: pass
    
    # Inspect 'Chi nhánh Fair Consulting Co., LTD Nhật Bản tại Việt Nam' in T2 C4,5
    try:
        inspect_for_hidden_chars(doc.tables[1].cell(3, 4).text, "T2 C4,5", f)
    except: pass
