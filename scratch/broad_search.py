
import docx
import sys
import io
import re

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def broad_search(text):
    text = text.upper()
    return "SỐ" in text and "DN" in text

try:
    doc = docx.Document("fs_C&C_2025_V.docx")
    print("Broad search in entire document...")
    # Body
    for i, p in enumerate(doc.paragraphs):
        if broad_search(p.text):
            print(f"Para[{i}]: '{p.text}'")
            
    # Tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if broad_search(cell.text):
                    print(f"Table[{ti}] Row[{ri}] Cell[{ci}]: '{cell.text}'")
                    
    # Headers
    for si, section in enumerate(doc.sections):
        for type_attr in ['header', 'footer', 'first_page_header', 'first_page_footer']:
            h = getattr(section, type_attr, None)
            if h:
                for pi, p in enumerate(h.paragraphs):
                    if broad_search(p.text):
                        print(f"Section[{si}] {type_attr} Para[{pi}]: '{p.text}'")
                for ti, table in enumerate(h.tables):
                    for ri, row in enumerate(table.rows):
                        for ci, cell in enumerate(row.cells):
                            if broad_search(cell.text):
                                print(f"Section[{si}] {type_attr} Table[{ti}] Row[{ri}] Cell[{ci}]: '{cell.text}'")
except Exception as e:
    print(f"Error: {e}")
