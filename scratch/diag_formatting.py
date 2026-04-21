
import docx
from docx.text.paragraph import Paragraph
import re
import translation_lib as tl

def test_on_real_doc(docx_path):
    print(f"Loading {docx_path}...")
    doc = docx.Document(docx_path)
    
    # Let's try to perform the swap manually on one known cell
    # Cell(4, 5) has '91.504.195'
    table = doc.tables[0]
    cell = table.cell(4, 5)
    print(f"DEBUG: Cell(4, 5) text before: '{cell.text}'")
    
    # We call apply_financial_number_formatting directly
    tl.apply_financial_number_formatting(doc, "E")
    
    print(f"DEBUG: Cell(4, 5) text after: '{cell.text}'")
    
    # If it still hasn't changed, let's see what runs are there NOW
    for p in cell.paragraphs:
        for r in p.runs:
            print(f"  Run text: '{r.text}'")

if __name__ == "__main__":
    test_on_real_doc("number.docx")
