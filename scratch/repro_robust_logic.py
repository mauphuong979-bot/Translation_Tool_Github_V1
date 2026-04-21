
import docx
import sys
import io
import os
import re
import unicodedata

# Import the library
import translation_lib as tl

# Set stdout to utf-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def broad_check(doc, label):
    print(f"\n--- Verification: {label} ---")
    
    # Check for Form Indicator Tabs
    print("\n[Form Indicator Check]")
    pattern = re.compile(r"(MẪU\s+SỐ|MÃ\s+SỐ|FORM|表\s*格).*?DN", re.IGNORECASE)
    for i, p in enumerate(doc.paragraphs):
        text = unicodedata.normalize('NFC', p.text)
        if pattern.search(text):
            has_tab = "\t" in p.text
            print(f"Para[{i}]: '{p.text[:50]}...' -> Tab present: {has_tab}")
            for r in p.runs:
                if pattern.search(unicodedata.normalize('NFC', r.text)):
                    print(f"  Run: '{r.text}' -> Bold: {r.bold}")

    # Check for Number Formatting in Tables
    print("\n[Number Formatting Check]")
    # We look for English style (dots as decimal)
    en_num_pattern = re.compile(r'\d+(?:,\d{3})*\.\d+')
    en_count = 0
    for table in doc.tables:
        for cell in table._cells:
            if en_num_pattern.search(cell.text):
                en_count += 1
    print(f"Found {en_count} English-style numbers in tables.")

# 1. Load the sample file
source_file = "fs_C&C_2025_V.docx"
if not os.path.exists(source_file):
    print(f"Error: {source_file} not found.")
    sys.exit(1)

doc = docx.Document(source_file)

# 2. Simulate processing
# We need a dummy translation map that includes the phrases from the doc
translation_map = {
    "BẢNG CÂN ĐỐI KẾ TOÁN": "BALANCE SHEET",
    "BÁO CÁO KẾT QUẢ HOẠT ĐỘNG KINH DOANH": "INCOME STATEMENT",
    "MẪU SỐ B 01 - DN": "FORM B 01 - DN",
    "MẪU SỐ B 02 - DN": "FORM B 02 - DN"
}

# Run the core replacement logic
tl.replace_text_in_document(
    doc, 
    translation_map, 
    case_threshold=30, 
    target_col="E",
    process_settings={
        "unicode": True,
        "dictionary": True,
        "number_swap": True,
        "date_format": False # Skip dates for now to focus on numbers/indicators
    }
)

# 3. Verify
broad_check(doc, "PROCESSED RESULT")

# Save for manual inspection if needed
output_file = "scratch/verified_output.docx"
doc.save(output_file)
print(f"\nSaved verified output to {output_file}")
