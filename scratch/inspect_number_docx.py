
import docx
from docx.text.paragraph import Paragraph

def inspect_numbers(docx_path):
    doc = docx.Document(docx_path)
    print(f"Total Tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i} ---")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # We check for paragraphs in cell
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text
                    if any(char.isdigit() for char in text):
                        print(f"Cell({r_idx}, {c_idx}) Para {p_idx}: '{text}'")
                        print(f"  Runs: {[run.text for run in para.runs]}")

if __name__ == "__main__":
    inspect_numbers("number.docx")
