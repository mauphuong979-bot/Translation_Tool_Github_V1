
import docx
import re

def inspect_doc(path, label):
    print(f"\n--- Inspecting {label}: {path} ---")
    doc = docx.Document(path)
    
    # 1. Check for Form Indicator Tabs
    print("\n[Form Indicator Check]")
    pattern = re.compile(r"(MẪU\s+SỐ|FORM|表\s*格).*?DN", re.IGNORECASE)
    found_form = False
    for para in doc.paragraphs:
        if pattern.search(para.text):
            found_form = True
            has_tab = "\t" in para.text
            print(f"Para: '{para.text[:50]}...' -> Tab present: {has_tab}")
            # Also check runs for bolding
            for run in para.runs:
                if pattern.search(run.text):
                    print(f"  Run: '{run.text}' -> Bold: {run.bold}")
                    
    if not found_form:
        # Check in tables (headers often in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if pattern.search(para.text):
                            found_form = True
                            has_tab = "\t" in para.text
                            print(f"TableCell Para: '{para.text[:50]}...' -> Tab present: {has_tab}")
                            for run in para.runs:
                                if pattern.search(run.text):
                                    print(f"  Run: '{run.text}' -> Bold: {run.bold}")

    # 2. Check for Number Formatting in Tables (especially merged ones)
    print("\n[Number Formatting Check]")
    # Look for numbers with commas as decimals (1.234,56 or 6,78) -> Vietnamese
    # Look for numbers with dots as decimals (1,234.56 or 6.78) -> English
    vn_num_pattern = re.compile(r'\d+(?:\.\d{3})*,\d+')
    en_num_pattern = re.compile(r'\d+(?:,\d{3})*\.\d+')
    
    vn_count = 0
    en_count = 0
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if vn_num_pattern.search(text): vn_count += 1
                if en_num_pattern.search(text): en_count += 1
                
    print(f"Found {vn_count} Vietnamese-style numbers.")
    print(f"Found {en_count} English-style numbers.")

print("Comparing Local vs Streamlit outputs...")
inspect_doc("fs_C&C_2025_V_E_tool_210426localhost.docx", "LOCAL")
inspect_doc("fs_C&C_2025_V_E_tool_210426_streamlit.docx", "STREAMLIT")
