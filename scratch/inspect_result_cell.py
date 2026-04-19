
from docx import Document
import os
import sys

# Reconfigure stdout to use UTF-8
if sys.stdout.encoding != 'utf-8':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'replace')
    except:
        pass

doc_path = r'd:\AI\Python\11_Project\Translation_Tool\Translation_Tool_Antigravity_V2\test_translated.docx'
if os.path.exists(doc_path):
    doc = Document(doc_path)
    cell = doc.tables[0].cell(12, 0)
    # Write repr to a file to be safe
    with open('scratch/inspect_cell_repr.txt', 'w', encoding='utf-8') as f:
        f.write(f"Cell text: {repr(cell.text)}\n")
        for p in cell.paragraphs:
            f.write(f"  Para text: {repr(p.text)}\n")
else:
    print("File not found")
