
import docx

def verify_fields(docx_path):
    doc = docx.Document(docx_path)
    # Check Cell(12, 2)
    table = doc.tables[0]
    cell = table.cell(12, 2)
    print(f"Cell(12, 2) text: '{cell.text}'")
    
    # Check XML for field markers
    p = cell.paragraphs[0]
    markers = p._element.xpath('.//w:fldChar | .//w:instrText')
    print(f"Found {len(markers)} field markers in Cell(12, 2).")
    for m in markers:
        tag = m.tag.split('}')[-1]
        print(f"  Marker: {tag}")

if __name__ == "__main__":
    verify_fields("scratch/number_output.docx")
