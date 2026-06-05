import sys
import os

# Add project directory to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from docx import Document

def main():
    docx_path = "../fs_AHVN_310326_V.docx"
    if not os.path.exists(docx_path):
        docx_path = "fs_AHVN_310326_V.docx"
    
    doc = Document(docx_path)
    print(f"Total tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        try:
            print(f"Table {i}: checking rows and columns...")
            print(f"  Rows count: {len(table.rows)}")
            print(f"  Columns count: {len(table.columns)}")
            # Try to access _cells
            cells = table._cells
            print(f"  _cells count: {len(cells)}")
        except Exception as e:
            print(f"  Table {i} failed on _cells: {type(e).__name__}: {e}")
            
            # Let's inspect rows/cells XML structure for this table
            try:
                print(f"  Let's see if we can iterate rows and cells manually:")
                for r_idx, row in enumerate(table.rows):
                    try:
                        row_cells = row.cells
                        print(f"    Row {r_idx} cell count: {len(row_cells)}")
                    except Exception as re:
                        print(f"    Row {r_idx} failed: {type(re).__name__}: {re}")
            except Exception as row_err:
                print(f"    Failed to iterate rows: {row_err}")

if __name__ == "__main__":
    main()
