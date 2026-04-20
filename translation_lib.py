import pandas as pd
import os
import re
import json
import unicodedata
import openpyxl
import difflib
from datetime import datetime
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Use absolute path for Streamlit Cloud compatibility
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

CLEANV_JSON = os.path.join(BASE_DIR, "clean_v.json")
CLEANV_XLSX = os.path.join(BASE_DIR, "CleanV.xlsx")

PARA_TEMPLATE_JSON = os.path.join(BASE_DIR, "para_template.json")
PARA_TEMPLATE_XLSX = os.path.join(BASE_DIR, "ParaTemplate.xlsx")

DICTIONARY_V3_XLSX = os.path.join(BASE_DIR, "Dictionary_v3.xlsx")
DICTIONARY_V3_JSON = os.path.join(BASE_DIR, "dictionary_v3.json")

# Legacy compatibility (optional, but keep if used elsewhere)
CLEANV_FILE = CLEANV_JSON
PARA_TEMPLATE_FILE = PARA_TEMPLATE_JSON
DICTIONARY_V3_FILE = DICTIONARY_V3_XLSX

def clean_text(text, preserve_newlines=False):
    """
    Ultra-robust text cleaning for Word documents.
    Normalizes to NFC, strips invisible controls/soft hyphens, and collapses all whitespace.
    This ensures matching works even with hidden formatting characters or different XML namespaces.
    """
    if not isinstance(text, str) or pd.isna(text):
        return ""
    # Normalize to NFC (Normalization Form C) to ensure consistent Vietnamese character encoding
    text = unicodedata.normalize('NFC', str(text))
    
    # Remove non-printable control characters, soft hyphens (\u00ad), and zero-width characters
    # \u200b: zero width space, \u200c: zero width non-joiner, \u200d: zero width joiner, \u2060: word joiner, \ufeff: BOM
    # \xb7: middle dot (·), \u2022: bullet point (•)
    # \u202a-\u202e: BiDi markers (LRE, RLE, PDF, LRO, RLO), \u200e: LTR mark, \u200f: RTL mark
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\xad\u200b\u200c\u200d\u2060\ufeff\xb7\u2022\u202a-\u202e\u200e\u200f]', ' ', text)
    
    if preserve_newlines:
        # Replace tabs and multiple spaces with a single space, but keep newlines
        text = re.sub(r'[ \t\r\f\v\u00a0]+', ' ', text)
    else:
        # Replace all whitespace sequences (tabs, newlines, non-breaking spaces) with a standard space
        text = re.sub(r'\s+', ' ', text)
    
    # Robustly handle numbering by ensuring exactly one space after dots at the start of strings or after spaces
    text = re.sub(r'(^|\s)(\d+)\.\s*', r'\1\2. ', text)
    
    return text.strip()

def ensure_proper_case(text):
    """
    If the text is all uppercase (including Vietnamese characters), 
    converts it to Title Case (Proper Case).
    """
    if not isinstance(text, str) or not text:
        return text
    
    # isupper() is True if all cased characters in S are uppercase
    if text.isupper():
        return text.title()
    return text

def load_cleanv_map():
    """
    Loads Vietnamese text normalization map from clean_v.json.
    """
    if os.path.exists(CLEANV_FILE):
        try:
            with open(CLEANV_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading clean_v.json: {e}")
            return {}
    return {}

def load_para_template_map():
    """
    Loads Paragraph Templates from para_template.json.
    """
    if os.path.exists(PARA_TEMPLATE_FILE):
        try:
            with open(PARA_TEMPLATE_FILE, 'r', encoding='utf-8') as f:
                raw_data = json.load(f)
                # Normalize keys for robust matching
                normalized_map = {}
                for vn_key, trans_data in raw_data.items():
                    normalized_map[clean_text(vn_key)] = trans_data
                return normalized_map
        except Exception as e:
            print(f"Error loading para_template.json: {e}")
            return {}
    return {}

def _apply_para_templates_to_container(container, para_map, target_col, replaced_paragraphs):
    """
    Helper to apply paragraph templates to a container (Document, Header, Footer).
    """
    count = 0
    # docx Paragraphs
    for para in container.paragraphs:
        # We clean the text for matching
        text = clean_text(para.text)
        if not text: continue
        
        # Sort keys by length descending
        sorted_keys = sorted(para_map.keys(), key=len, reverse=True)
        para_text_lower = text.lower()
        
        for vn_key in sorted_keys:
            # Case-insensitive check
            if vn_key.lower() in para_text_lower:
                # Found a template match
                template_data = para_map[vn_key]
                new_text = template_data.get(target_col, "")
                if new_text:
                    # print(f"DEBUG: Found ParaTemplate match! Key: '{vn_key[:30]}...' -> Target: '{target_col}'")
                    # Replace entire paragraph text while preserving the best available formatting
                    # We use the full paragraph text as the key to ensure the whole paragraph is swapped.
                    full_original = clean_text(para.text)
                    dummy_map = prepare_translation_list({full_original: new_text})
                    apply_translations_to_paragraph(para, dummy_map)
                    replaced_paragraphs.add(para) # Mark as replaced
                    count += 1
                    break # Only one template per paragraph
                    
    return count

def apply_paragraph_templates(doc, para_map, target_col):
    """
    Applies paragraph-level templates to the entire document.
    Returns the count of replacements and a set of modified paragraph objects.
    """
    if not para_map:
        return 0, set()
        
    total_count = 0
    replaced_paragraphs = set()

    # 1. Main body
    total_count += _apply_para_templates_to_container(doc, para_map, target_col, replaced_paragraphs)
    
    # 2. Tables (Cells are containers)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_count += _apply_para_templates_to_container(cell, para_map, target_col, replaced_paragraphs)
                
    # 3. Headers and Footers
    for section in doc.sections:
        total_count += _apply_para_templates_to_container(section.header, para_map, target_col, replaced_paragraphs)
        total_count += _apply_para_templates_to_container(section.footer, para_map, target_col, replaced_paragraphs)
        if section.different_first_page_header_footer:
            total_count += _apply_para_templates_to_container(section.first_page_header, para_map, target_col, replaced_paragraphs)
            total_count += _apply_para_templates_to_container(section.first_page_footer, para_map, target_col, replaced_paragraphs)
        try:
            total_count += _apply_para_templates_to_container(section.even_page_header, para_map, target_col, replaced_paragraphs)
            total_count += _apply_para_templates_to_container(section.even_page_footer, para_map, target_col, replaced_paragraphs)
        except: pass
        
    return total_count, replaced_paragraphs

def _apply_cleanv_to_container(container, cleanv_map):
    """
    Helper to apply CleanV normalization (substring replacement) to all text in a container.
    """
    count = 0
    if not cleanv_map: return 0
    
    # Sort keys by length descending
    sorted_keys = sorted(cleanv_map.keys(), key=len, reverse=True)
    
    for para in container.paragraphs:
        changed_para = False
        for run in para.runs:
            if not run.text: continue
            new_text, changed = apply_normalization_to_text(run.text, cleanv_map)
            if changed:
                run.text = new_text
                changed_para = True
        if changed_para:
            count += 1
            
    return count

def apply_unicode_normalization(doc):
    """
    Explicitly normalizes all text in the document (Body, Header, Footer, Tables) to NFC.
    """
    def _norm_container(container):
        for para in container.paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = unicodedata.normalize('NFC', run.text)

    # 1. Body
    _norm_container(doc)
    
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _norm_container(cell)
                
    # 3. Headers and Footers
    for section in doc.sections:
        headers = [section.header, section.footer]
        if section.different_first_page_header_footer:
            headers.extend([section.first_page_header, section.first_page_footer])
        try:
            headers.extend([section.even_page_header, section.even_page_footer])
        except: pass
        
        for h in headers:
            if h:
                _norm_container(h)
                for table in h.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            _norm_container(cell)

def apply_cleanv_normalization(doc, cleanv_map):
    """
    Applies Vietnamese text normalization to the entire document.
    """
    if not cleanv_map:
        return 0
        
    total_count = 0
    # 1. Main body
    total_count += _apply_cleanv_to_container(doc, cleanv_map)
    
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_count += _apply_cleanv_to_container(cell, cleanv_map)
                
    # 3. Headers and Footers
    for section in doc.sections:
        total_count += _apply_cleanv_to_container(section.header, cleanv_map)
        total_count += _apply_cleanv_to_container(section.footer, cleanv_map)
        if section.different_first_page_header_footer:
            total_count += _apply_cleanv_to_container(section.first_page_header, cleanv_map)
            total_count += _apply_cleanv_to_container(section.first_page_footer, cleanv_map)
        try:
            total_count += _apply_cleanv_to_container(section.even_page_header, cleanv_map)
            total_count += _apply_cleanv_to_container(section.even_page_footer, cleanv_map)
        except: pass
        
    return total_count

def apply_special_textbox_formatting(doc, target_col):
    """
    Finds "BẢN DỰ THẢO" in textboxes and replaces it with language-specific text.
    Also applies special formatting (10.5pt, margins, autosize).
    """
    find_text = "BẢN DỰ THẢO"
    
    replace_data = {
        "E": "DRAFT",
        "Hs": "草 稿",
        "Ht": "草 稿"
    }
    
    target_text = replace_data.get(target_col, find_text)
    is_chinese = target_col in ["Hs", "Ht"]
    
    # 1. Find all textbox contents in body, headers, footers
    # We use a combined list of all XML elements for a thorough hunt
    elements = [doc._element]
    for section in doc.sections:
        elements.extend([section.header._element, section.footer._element])
        if section.different_first_page_header_footer:
            elements.extend([section.first_page_header._element, section.first_page_footer._element])
        try:
            elements.extend([section.even_page_header._element, section.even_page_footer._element])
        except: pass

    count = 0
    for root in elements:
        textboxes = root.xpath('.//*[local-name()="txbxContent"]')
        for txbx in textboxes:
            # Check if it contains the target draft phrase
            txt_nodes = txbx.xpath('.//*[local-name()="t"]')
            full_text = "".join(t.text for t in txt_nodes if t.text)
            
            if find_text in full_text:
                # 2. Perform Replacement and Font Styling
                p_nodes = txbx.xpath('.//*[local-name()="p"]')
                for p_node in p_nodes:
                    para = Paragraph(p_node, doc)
                    if find_text in para.text:
                        para.text = para.text.replace(find_text, target_text)
                    
                    # Force font styling to all runs
                    for run in para.runs:
                        if is_chinese:
                            _set_run_fonts_refined(run, "Times New Roman", "DFKai-SB", 10.5)
                        else:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(10.5)
                
                # 3. XML Level Formatting (Margins and AutoSize)
                try:
                    # VBA MarginLeft=0, MarginRight=0, AutoSize=True, WordWrap=True
                    # VML Style (older textboxes)
                    v_textbox = txbx.getparent()
                    if v_textbox is not None and v_textbox.tag.endswith('textbox'):
                        v_textbox.set('inset', '0,0,0,0') # 0 margins
                        
                        v_shape = v_textbox.getparent()
                        if v_shape is not None and v_shape.tag.endswith('shape'):
                            style = v_shape.get('style', '')
                            # Add mso-fit-shape-to-text:t for AutoSize
                            if 'mso-fit-shape-to-text' not in style:
                                style += ';mso-fit-shape-to-text:t'
                            v_shape.set('style', style)
                    
                    # DrawingML Style (modern textboxes)
                    # We look for a sister element <wps:bodyPr> which controls margins/wrap
                    # txbxContent -> txbx -> (sister) bodyPr
                    wps_txbx = txbx.getparent()
                    if wps_txbx is not None and wps_txbx.tag.endswith('txbx'):
                        wps_wsp = wps_txbx.getparent()
                        if wps_wsp is not None:
                            bodyPr = wps_wsp.find(qn('wps:bodyPr'))
                            if bodyPr is not None:
                                bodyPr.set('lIns', '0')
                                bodyPr.set('rIns', '0')
                                bodyPr.set('tIns', '0')
                                bodyPr.set('bIns', '0')
                                # AutoSize in modern Word is often complex, but 0 margins helps.
                except:
                    pass
                count += 1
    return count

def apply_sizing_and_layout(doc, target_col="E"):
    """
    Applies specialized table sizing and layout formatting.
    """
    if not doc.tables:
        return

    # 1. SPECIAL CASE: Table 1 (Cover)
    # VBA logic: tbl.Columns.Count = 1 And tbl.Rows.Count = 7
    try:
        t1 = doc.tables[0]
        if len(t1.columns) == 1 and len(t1.rows) == 7:
            if target_col == "E":
                # VBA indices are 1-based. Row 1 is index 0.
                # R1: 14, R2: 10.5, R3: 10.5, R6: 12, R7: 12
                for i, size in [(0, 14), (1, 10.5), (2, 10.5), (5, 12), (6, 12)]:
                    for para in t1.rows[i].cells[0].paragraphs:
                        if not para.runs: para.add_run("") # Ensure at least one run
                        for run in para.runs:
                            run.font.size = Pt(size)
            elif target_col in ["Hs", "Ht"]:
                # R1: 16, R2: 13, R3: 10, R6: 15, R7: 15
                for i, size in [(0, 16), (1, 13), (2, 10), (5, 15), (6, 15)]:
                    for para in t1.rows[i].cells[0].paragraphs:
                        if not para.runs: para.add_run("")
                        for run in para.runs:
                            run.font.size = Pt(size)
    except: pass

    # 2. GLOBAL TABLE LOGIC
    for tbl in doc.tables:
        try:
            row_count = len(tbl.rows)
            col_count = len(tbl.columns)
            
            # Row count 37-40 -> Row 1 height = 1.01 cm
            if 37 <= row_count <= 40:
                tbl.rows[0].height = Cm(1.01)
                
            # Cols > 10 -> Font size = 7pt for all cells
            if col_count > 10:
                for row in tbl.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(7)
                                
            # 3-column table with specific text in Cell(0, 2)
            if col_count == 3:
                # VBA: tbl.cell(1, 3).Range.text (1-indexed) -> Cell(0, 2)
                cell_text = tbl.cell(0, 2).text
                if "Percentage of interest (%)" in cell_text:
                    # Setting width at cell level per row for better compatibility
                    for row in tbl.rows:
                        row.cells[1].width = Cm(4.98)
                        row.cells[2].width = Cm(4.07)
        except: pass

def _set_run_fonts_refined(run, latin_font, east_asia_font, size_pt=None):
    """
    Sets dual font families and optional size.
    Strengthened for maximum compatibility by setting all 4 font slots.
    """
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    
    # Set all slots to ensure consistent rendering across locales
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rFonts.set(qn('w:cs'), latin_font) # Complex scripts slots
    
    # Always hint how to handle the run
    if contains_chinese(run.text):
        rFonts.set(qn('w:hint'), 'eastAsia')
    else:
        rFonts.set(qn('w:hint'), 'default')
    
    if size_pt:
        run.font.size = Pt(size_pt)

def _process_paragraph_font_dual(para, target_col):
    """
    Splits runs in a paragraph to handle mixed Chinese/Latin formatting.
    Chinese -> DFKai-SB, 10pt
    Latin -> Times New Roman, Original Size
    """
    if not para.text:
        return

    old_runs = list(para.runs)
    if not old_runs:
        return

    # Clear existing runs
    p_el = para._element
    
    LATIN_FONT = "Times New Roman"
    CJK_FONT = "DFKai-SB"

    # NEW: Protection for Link Fields (Excel links, automated data)
    # If fields are present, we DO NOT split runs because it destroys the XML structure of the links.
    if has_fields(para):
        for run in para.runs:
            if not run.text: continue
            # Non-destructive formatting: apply font/size to existing runs
            is_ch = contains_chinese(run.text)
            target_size = 10 if is_ch else None # Enforce 10pt for Chinese result text
            _set_run_fonts_refined(run, LATIN_FONT, CJK_FONT, target_size)
        return

    for r in old_runs:
        p_el.remove(r._element)
        
    # Regex to split by Chinese characters while keeping them
    cjk_pattern = r'([\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]+)'
    
    for old_run in old_runs:
        text = old_run.text
        if not text: continue
        
        # Get original size to preserve for Latin text
        orig_size_pt = None
        try:
            if old_run.font and old_run.font.size:
                orig_size_pt = old_run.font.size.pt
        except: pass
        
        if not contains_chinese(text):
            # Pure Latin/Numbers -> Keep original size
            new_run = para.add_run(text)
            _copy_run_format(old_run, new_run)
            _set_run_fonts_refined(new_run, LATIN_FONT, CJK_FONT, orig_size_pt)
        else:
            # Mixed content -> Split into small chunks
            parts = re.split(cjk_pattern, text)
            for part in parts:
                if not part: continue
                new_run = para.add_run(part)
                _copy_run_format(old_run, new_run)
                
                if contains_chinese(part):
                    # Chinese segment -> Enforce 10pt and CJK font
                    _set_run_fonts_refined(new_run, LATIN_FONT, CJK_FONT, 10)
                else:
                    # Latin/Number segment -> Keep original size and TNR font
                    _set_run_fonts_refined(new_run, LATIN_FONT, CJK_FONT, orig_size_pt)

def apply_chinese_font_formatting(doc, target_col):
    """
    Applies Chinese report font standards with dual-font and size preservation.
    """
    if target_col not in ["Hs", "Ht"]:
        return

    # 1. Main body
    for para in doc.paragraphs:
        _process_paragraph_font_dual(para, target_col)
    
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph_font_dual(para, target_col)
                
    # 3. Headers and Footers
    for section in doc.sections:
        # Loop through all possible headers/footers
        headers = [section.header, section.footer]
        if section.different_first_page_header_footer:
            headers.extend([section.first_page_header, section.first_page_footer])
        try:
            headers.extend([section.even_page_header, section.even_page_footer])
        except: pass
        
        for h in headers:
            if h:
                for para in h.paragraphs:
                    _process_paragraph_font_dual(para, target_col)
                for table in h.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _process_paragraph_font_dual(para, target_col)

def format_dates_in_tables(doc, target_col="E"):
    """
    Finds dates in DD/MM/YYYY format in table cells and reformats them.
    English (E): MMM DD, YYYY
    Chinese (Hs/Ht): YYYY/MM/DD日
    
    Now works at the run level to preserve Word fields (Excel links).
    """
    # Regex for DD/MM/YYYY (supports 1 or 2 digits for day/month)
    date_pattern = re.compile(r'^(0?[1-9]|[12][0-9]|3[01])/(0?[1-9]|1[0-2])/([0-9]{4})$')
    
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Iterate through all paragraphs and runs inside the cell
                # This allows us to find and replace dates even if they are part of a field
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text:
                            continue
                            
                        # Aggressive cleaning: remove control chars < 32 and trim
                        clean_run_text = "".join(ch for ch in run.text if ord(ch) >= 32).strip()
                        
                        # Check if run text matches EXACTLY a date
                        match = date_pattern.match(clean_run_text)
                        if match:
                            try:
                                # Parse date
                                day, month, year = match.groups()
                                dt = datetime(int(year), int(month), int(day))
                                
                                # Reformat based on target language
                                if target_col == "E":
                                    # 'MMM dd, yyyy' -> 'Jan 01, 2023'
                                    new_text = dt.strftime("%b %d, %Y")
                                elif target_col in ["Hs", "Ht"]:
                                    # 'yyyy/mm/dd' & "日" -> '2023/12/31日'
                                    new_text = dt.strftime("%Y/%m/%d") + "日"
                                else:
                                    continue
                                    
                                # Preserve any extra characters that were trimmed
                                run.text = run.text.replace(clean_run_text, new_text)
                                count += 1
                                
                                # NEW: Disconnect Excel fields for this cell
                                # This ensures the translated date is static text.
                                unlink_fields_in_item(cell)
                                
                                # Since we've modified the XML structure of the cell by unlinking,
                                # it's safer to stop iterating THIS cell and move to the next.
                                break # break para loop
                            except:
                                pass
                    else: continue
                    break # break cell loop (para loop was broken)
    return count

def contains_vietnamese(text):
    """
    Checks if a string contains characters unique to the Vietnamese language.
    Includes NFC-normalized characters.
    """
    if not text:
        return False
        
    # Standard list of specific Vietnamese characters (NFC)
    # Includes lowercase and uppercase variants
    vn_chars = "àáảãạâầấẩẫậăằắẳẵặèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ"
    vn_chars += vn_chars.upper()
    vn_set = set(vn_chars)
    
    for char in text:
        if char in vn_set:
            return True
    return False

def remove_accents(text):
    """
    Removes Vietnamese diacritics from text.
    'Nguyễn Văn A' -> 'Nguyen Van A'
    """
    if not isinstance(text, str) or not text:
        return text
    
    # Handle đ/Đ manually as they are not base characters in NFC/NFD
    text = text.replace('đ', 'd').replace('Đ', 'D')
    
    # Normalize with NFD (Normalization Form Decomposition) to separate base chars from marks
    nfd_text = unicodedata.normalize('NFD', text)
    # Filter out characters that are in the 'Mn' (Mark, Nonspacing) category
    result = "".join(c for c in nfd_text if unicodedata.category(c) != 'Mn')
    
    return unicodedata.normalize('NFC', result)

def apply_signer_accent_removal(doc, metadata):
    """
    Replaces accented signer names with unaccented versions throughout the document.
    Handles field links by unlinking then replacing.
    """
    if not metadata:
        return 0
        
    signer_keys = ["signer_1", "signer_2", "signer_3"]
    replacements = {}
    
    for key in signer_keys:
        val = metadata.get(key)
        if val and contains_vietnamese(val):
            unaccented = remove_accents(val)
            if unaccented != val:
                replacements[val] = unaccented
                
    if not replacements:
        return 0
        
    count = 0
    # Process all containers including headers/footers
    from docx.text.paragraph import Paragraph
    
    # Use a helper to get all items (paragraphs and table cells)
    def _get_items():
        # Body
        for para in doc.paragraphs: yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: yield para
        # Headers/Footers
        for section in doc.sections:
            headers = [section.header, section.footer]
            if section.different_first_page_header_footer:
                headers.extend([section.first_page_header, section.first_page_footer])
            try: headers.extend([section.even_page_header, section.even_page_footer])
            except: pass
            for h in headers:
                if not h: continue
                for para in h.paragraphs: yield para
                for table in h.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs: yield para

    for para in _get_items():
        para_text = para.text
        found_signer = False
        for accented in replacements:
            if accented in para_text:
                found_signer = True
                break
        
        if found_signer:
            # Unlink fields if any to handle Excel links
            if has_fields(para):
                unlink_fields_in_item(para)
            
            # Perform replacement in runs
            for accented, unaccented in replacements.items():
                # We do a case-insensitive check but preserve case if possible?
                # The user said "nếu là tiếng Việt có dấu, thì thực hiện thay thế bằng tiếng Việt không dấu"
                # We'll use case-insensitive regex for finding but careful replacement
                pattern = re.compile(re.escape(accented), re.IGNORECASE)
                
                # Scan runs
                for run in para.runs:
                    if not run.text: continue
                    new_text = pattern.sub(unaccented, run.text)
                    if new_text != run.text:
                        run.text = new_text
                        count += 1
                        
    return count

def contains_chinese(text):
    """
    Checks if a string contains CJK characters (Chinese characters) 
    including common CJK symbols and punctuation.
    """
    if not text:
        return False
    pattern = r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff\u3000-\u303f\uff00-\uffef]'
    return bool(re.search(pattern, text))

def swap_vn_to_en_number_separators(text):
    """
    Swaps Vietnamese number separators (. for thousands, , for decimal)
    to International/English style (, for thousands, . for decimal).
    
    Target patterns:
    - 1.234,56 -> 1,234.56
    - 1.234 -> 1,234 (Must have exactly 3 digits after dot)
    - 6,78 -> 6.78
    
    Safe from dates like 31.12.2025 because those don't have 3-digit groups.
    """
    if not text:
        return text
        
    # Pattern explanation:
    # (?<!\d) : No digit before
    # (?: ... ) : Non-capturing alternatives
    #   \d{1,3}(?:\.\d{3})+(?:,\d+)? : Numbers with thousands dots and optional decimal comma
    #   | \d+,\d+ : Numbers with only decimal comma
    # (?!\d) : No digit after
    pattern = re.compile(r'(?<!\d)(?:\d{1,3}(?:\.\d{3})+(?:,\d+)?|\d+,\d+)(?!\d)')
    
    def replace_func(match):
        val = match.group(0)
        # 3-step swap to avoid overwriting
        return val.replace('.', 'TEMP_DOT').replace(',', '.').replace('TEMP_DOT', ',')
        
    return pattern.sub(replace_func, text)

def apply_financial_number_formatting(doc, target_col):
    """
    Safe run-level scan to swap number separators.
    Used for English and Chinese reports to match international standards.
    Operates ONLY on run.text to preserve Link Fields (Excel links).
    """
    if target_col == "V":
        return

    # Process all possible containers
    containers = [doc]
    for section in doc.sections:
        containers.extend([section.header, section.footer])
        if section.different_first_page_header_footer:
            containers.extend([section.first_page_header, section.first_page_footer])
        try:
            containers.extend([section.even_page_header, section.even_page_footer])
        except: pass

    for container in containers:
        # 1. Body Paragraphs
        for para in container.paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = swap_vn_to_en_number_separators(run.text)
        
        # 2. Table Cells
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Deep traversal of cell contents
                    for p_el in cell._element.xpath('.//*[local-name()="p"]'):
                        para = Paragraph(p_el, cell)
                        for run in para.runs:
                            if run.text:
                                run.text = swap_vn_to_en_number_separators(run.text)
                    
                    # Handle text boxes within tables too
                    for t_el in cell._element.xpath('.//*[local-name()="textbox"]'):
                        for p_el in t_el.xpath('.//*[local-name()="p"]'):
                            para = Paragraph(p_el, t_el)
                            for run in para.runs:
                                if run.text:
                                    run.text = swap_vn_to_en_number_separators(run.text)

def has_fields(doc_item):
    """
    Checks if a paragraph or cell contains Word fields (links to Excel, automated data).
    Detecting w:fldSimple, w:fldChar, or w:instrText ensures we don't destroy link structures.
    """
    try:
        element = doc_item._element
        # XPath to find common field markers
        fields = element.xpath('.//*[local-name()="fldSimple" or local-name()="fldChar" or local-name()="instrText"]')
        return len(fields) > 0
    except:
        return False

def unlink_fields_in_item(item):
    """
    Strips Word fields (fldSimple, complex fields) to leave only plain text results (static text).
    This "disconnects" the content from external sources like Excel.
    """
    try:
        element = item._element
        
        # 1. Handle fldSimple: Replace the field with its children content
        for fld in element.xpath('.//*[local-name()="fldSimple"]'):
            parent = fld.getparent()
            if parent is not None:
                # Move all children (runs) out of the fldSimple element
                for child in list(fld):
                    fld.addprevious(child)
                parent.remove(fld)
                
        # 2. Handle complex fields (begin, separate, end markers and instrText)
        # We remove the runs that contain these markers, leaving only the "result" runs.
        # Field markers are usually wrapped in <w:r>
        xpath_query = './/*[local-name()="fldChar" or local-name()="instrText"]'
        for marker in element.xpath(xpath_query):
            # Find the parent run element <w:r>
            run_element = marker.getparent()
            # If marker is instrText, its parent is <w:r>. 
            # If marker is fldChar, its parent is <w:r>.
            # However, sometimes they might be buried deeper if there are nested elements, 
            # but usually it's direct.
            while run_element is not None and not run_element.tag.endswith('}r'):
                run_element = run_element.getparent()
                
            if run_element is not None:
                run_parent = run_element.getparent()
                if run_parent is not None:
                    run_parent.remove(run_element)
        return True
    except Exception as e:
        # Silently fail if XML manipulation hits an edge case
        return False

def _copy_run_format(src_run, dest_run):
    """
    Clones standard font formatting from src_run to dest_run.
    """
    dest_run.bold = src_run.bold
    dest_run.italic = src_run.italic
    dest_run.underline = src_run.underline
    if src_run.font.name:
        dest_run.font.name = src_run.font.name
    if src_run.font.size:
        dest_run.font.size = src_run.font.size
    try:
        if src_run.font.color and src_run.font.color.rgb:
            dest_run.font.color.rgb = src_run.font.color.rgb
    except: pass

def find_fuzzy_translation(text, translation_map, threshold=0.6):
    """
    Finds the most similar Vietnamese term in the dictionary and returns its translation.
    Uses difflib.get_close_matches for robust fuzzy matching.
    """
    if not text or not translation_map:
        return None
        
    # Standardize input for better matching
    clean_input = clean_text(text)
    if len(clean_input) < 2: return None
    
    keys = list(translation_map.keys())
    matches = difflib.get_close_matches(clean_input, keys, n=1, cutoff=threshold)
    
    if matches:
        return translation_map[matches[0]]
    return None

def _process_item_for_word_highlight(para):
    """
    Splits runs in a paragraph to highlight only the specific words containing Vietnamese.
    Returns (vn_total_length, vn_segments_list) for suggestion logic.
    """
    vn_total_len = 0
    vn_segments = set()
    
    old_runs = list(para.runs)
    if not old_runs:
        return 0, []

    # 1. Check if the paragraph has ANY Vietnamese first to avoid unnecessary splitting
    if not contains_vietnamese(para.text):
        return 0, []

    # 2. Clear existing runs from paragraph element safely
    p_el = para._element
    
    # NEW: Protection for Link Fields in Highlight step
    if has_fields(para):
        # If paragraph has links, only highlight existing runs without splitting
        for run in para.runs:
            if contains_vietnamese(run.text):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                vn_total_len += len(run.text)
                vn_segments.add(run.text.strip())
        return vn_total_len, list(vn_segments)

    for r in old_runs:
        p_el.remove(r._element)
    
    # 3. Define Vietnamese characters for regex
    vn_chars = "àáảãạâầấẩẫậăằắẳẵặèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ"
    vn_chars += vn_chars.upper()
    
    # Pattern to match "words": sequences of alphanumeric + VN chars
    # Splits by everything else (punctuation, space, etc.)
    pattern = r'([a-zA-Z0-9' + vn_chars + r']+)'
    
    for old_run in old_runs:
        text = old_run.text
        if not text: continue
        
        if not contains_vietnamese(text):
            # Simple run, just restore it
            new_run = para.add_run(text)
            _copy_run_format(old_run, new_run)
        else:
            # Run contains Vietnamese, need to split it
            parts = re.split(pattern, text)
            for part in parts:
                if not part: continue
                new_run = para.add_run(part)
                _copy_run_format(old_run, new_run)
                if contains_vietnamese(part):
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    vn_total_len += len(part)
                    vn_segments.add(part.strip())
    
    return vn_total_len, list(vn_segments)

def highlight_vietnamese_text(doc, translation_map=None, original_texts=None, show_suggestions=True):
    """
    Scans the entire document and highlights only the specific WORDS 
    containing Vietnamese text in yellow.
    If show_suggestions is True and a paragraph has > 30 Vietnamese characters, 
    adds fuzzy suggestions based on the similarity of the ORIGINAL paragraph text.
    """
    
    def handle_suggestions(para, vn_count):
        # Only process if toggled ON, highlighted VN characters > 30 and dictionary is available
        if show_suggestions and vn_count > 30 and translation_map:
            # Use original text if provided, otherwise fallback to current text
            # Use _element as key because id(para) is not stable across wrapper recreations
            raw_orig_text = original_texts.get(para._element, para.text) if original_texts else para.text
            
            # Clean and normalize for matching
            match_text = clean_text(raw_orig_text)
            if not match_text: return
            
            # Find the single best fuzzy match for the ORIGINAL text
            keys = list(translation_map.keys())
            matches = difflib.get_close_matches(match_text, keys, n=1, cutoff=0.6)
            
            if matches:
                best_match_key = matches[0]
                suggestion = translation_map[best_match_key]
                
                # Create combined text: Original VN + Suggested Translation
                # Using \n for newline within the same professional blue block
                combined_text = f"[Original: {raw_orig_text}]\n[Suggest: {suggestion}]"
                
                # Insert new paragraph IMMEDIATELY after the current one
                new_p_el = OxmlElement('w:p')
                para._element.addnext(new_p_el)
                new_p = Paragraph(new_p_el, para._parent)
                
                # Add text and format (Blue, Regular)
                run = new_p.add_run(combined_text)
                run.font.color.rgb = RGBColor(0, 51, 204) # Professional Blue
                return True
        return False

    # 1. Main body paragraphs
    # Iterating over list(doc.paragraphs) to allow modification (adding new paragraphs)
    for para in list(doc.paragraphs):
        count, _ = _process_item_for_word_highlight(para)
        handle_suggestions(para, count)
        
    # 2. Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in list(cell.paragraphs):
                    count, _ = _process_item_for_word_highlight(para)
                    handle_suggestions(para, count)
                    
    # 3. Headers and Footers
    for section in doc.sections:
        containers = [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer
        ]
        try:
            containers.extend([section.even_page_header, section.even_page_footer])
        except: pass
        
        for container in containers:
            if container:
                for para in list(container.paragraphs):
                    count, _ = _process_item_for_word_highlight(para)
                    handle_suggestions(para, count)
                # Tables in headers/footers
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in list(cell.paragraphs):
                                count, _ = _process_item_for_word_highlight(para)
                                handle_suggestions(para, count)

def apply_chinese_currency_cleanup(doc):
    """
    Final post-processing to fix specific currency redundant translations in Chinese.
    """
    corrections = {
        "越南盾（越南盾）": "越南盾（VND）",
        "美元（美元）": "美元（USD）"
    }
    
    def _clean_container(container):
        for para in container.paragraphs:
            original_text = para.text
            new_text = original_text
            for key, val in corrections.items():
                if key in new_text:
                    new_text = new_text.replace(key, val)
            if new_text != original_text:
                para.text = new_text

    # 1. Main body
    _clean_container(doc)
    
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _clean_container(cell)
def apply_normalization_to_text(text, cleanv_map):
    """
    Applies Vietnamese text normalization using the cleanv_map.
    This replaces "dirty" Vietnamese patterns with "cleaned" ones.
    """
    if not text or not cleanv_map:
        return text, False
    
    new_text = text
    changed = False
    
    # Iterate through keys in their original order (top-to-bottom in JSON)
    for key in cleanv_map:
        if key in new_text:
            new_text = new_text.replace(key, cleanv_map[key])
            changed = True
            
    return new_text, changed

def parse_date_to_tags(date_str, prefix):
    """
    Extracts day, month, year from various date formats (DD/MM/YYYY, DD-MM-YYYY, DD.MM.YYYY).
    Supports 2-digit and 4-digit years.
    Returns tags like [prefix_day], [prefix_month], [prefix_year].
    """
    if not date_str or not isinstance(date_str, str):
        return {f"[{prefix}_day]": "", f"[{prefix}_month]": "", f"[{prefix}_year]": ""}
    
    # Clean and split by common separators
    clean_date = date_str.strip()
    parts = re.split(r"[/.-]", clean_date)
    
    if len(parts) == 3:
        try:
            d, m, y = parts[0].strip(), parts[1].strip(), parts[2].strip()
            # Pad with zero if single digit
            if d.isdigit(): d = f"{int(d):02d}"
            if m.isdigit(): m = f"{int(m):02d}"
            # Handle 2-digit years if necessary (simple heuristic)
            if y.isdigit() and len(y) == 2:
                y = "20" + y # Assume 20xx
            
            return {f"[{prefix}_day]": d, f"[{prefix}_month]": m, f"[{prefix}_year]": y}
        except: pass
    return {f"[{prefix}_day]": "", f"[{prefix}_month]": "", f"[{prefix}_year]": ""}

def parse_period_to_tags(period_str, date1_prefix, date2_prefix):
    """
    Extracts two dates from a period string and maps to tags.
    Supports /, -, . as separators and 2/4-digit years.
    """
    tags = {}
    # Initialize default empty tags
    for p in [date1_prefix, date2_prefix]:
        tags.update({f"[{p}_day]": "", f"[{p}_month]": "", f"[{p}_year]": ""})
        
    if not period_str or not isinstance(period_str, str):
        return tags
        
    # Find all date patterns: \d{1,2} [separator] \d{1,2} [separator] \d{2,4}
    # Using re.findall returns tuples of groups
    dates = re.findall(r"(\d{1,2})[/.-](\d{1,2})[/.-](\d{2,4})", period_str)
    
    # Process First Date
    if len(dates) >= 1:
        d, m, y = dates[0]
        if len(y) == 2: y = "20" + y
        tags.update({
            f"[{date1_prefix}_day]": f"{int(d):02d}", 
            f"[{date1_prefix}_month]": f"{int(m):02d}", 
            f"[{date1_prefix}_year]": y
        })
        
    # Process Second Date
    if len(dates) >= 2:
        d, m, y = dates[1]
        if len(y) == 2: y = "20" + y
        tags.update({
            f"[{date2_prefix}_day]": f"{int(d):02d}", 
            f"[{date2_prefix}_month]": f"{int(m):02d}", 
            f"[{date2_prefix}_year]": y
        })
        
    return tags

def get_metadata_substitution_map(metadata):
    """
    Common mapping for both template filling and legacy replacement.
    Includes Vietnamese variants and supports v_name/V_NAME distinction.
    """
    name_vn = metadata.get("name_vn", "")
    name_trans = metadata.get("name_trans", "")
    report_date = metadata.get("report_date", "")
    
    # 1. Base Company Name Casing
    subs = {
        "[v_name]": name_vn,
        "[V_NAME]": name_vn.upper(),
        "[t_name]": name_trans,
        "[T_NAME]": name_trans.upper(),
        # Vietnamese Variants
        "[têncôngty]": name_vn,
        "[TÊN CÔNG TY]": name_vn.upper(),
        "[tên khách hàng]": name_trans,
        "[TÊN KHÁCH HÀNG]": name_trans.upper(),
        "[v_signer_1]": metadata.get("signer_1", ""),
        "[v_signer_2]": metadata.get("signer_2", ""),
        "[v_signer_3]": metadata.get("signer_3", ""),
        "[signer1]": metadata.get("signer_1", ""),
        "[signer2]": metadata.get("signer_2", ""),
        "[signer3]": metadata.get("signer_3", ""),
        "[ngườiký1]": metadata.get("signer_1", ""),
        "[ngườiký2]": metadata.get("signer_2", ""),
        "[ngườiký3]": metadata.get("signer_3", ""),
    }
    
    # Fallback/Legacy [v_signer]
    all_signers = []
    for i in range(1, 4):
        s = metadata.get(f"signer_{i}")
        if s: all_signers.append(s)
    subs["[v_signer]"] = ", ".join(all_signers)
    subs["[ngườiký]"] = ", ".join(all_signers)
    
    # 2. Date Tags (We add both casings because the output is numeric/consistent)
    year_end_tags = parse_date_to_tags(metadata.get("year_end", ""), "e")
    report_date_tags = parse_date_to_tags(report_date, "r")
    p1_tags = parse_period_to_tags(metadata.get("period_in", ""), "p1", "p2")
    p2_tags = parse_period_to_tags(metadata.get("period_in_2", ""), "p3", "p4")
    
    for tag_set in [year_end_tags, report_date_tags, p1_tags, p2_tags]:
        for k, v in tag_set.items():
            subs[k.lower()] = v
            subs[k.upper()] = v
            
    # Add other Vietnamese utility tags (both casings)
    utility_v = {
        "[ngàykếtthúcnăm]": metadata.get("year_end", ""),
        "[ngàybáocáo]": report_date,
    }
    for k, v in utility_v.items():
        subs[k.lower()] = v
        subs[k.upper()] = v

    # Handle [nămbáocáo] specifically
    year_only = ""
    if report_date and isinstance(report_date, str):
        parts = re.split(r"[/.-]", report_date)
        if len(parts) == 3:
            year_only = parts[2].strip()
    subs["[nămbáocáo]"] = year_only
    subs["[NĂM BÁO CÁO]"] = year_only
    
    # Normalize all keys to NFC
    return {unicodedata.normalize('NFC', k): v for k, v in subs.items()}

def sync_clean_v():
    """Syncs CleanV.xlsx to clean_v.json"""
    if not os.path.exists(CLEANV_XLSX):
        return False, f"Missing {CLEANV_XLSX}"
    try:
        df = pd.read_excel(CLEANV_XLSX).astype(str)
        result = {}
        for _, row in df.iterrows():
            orig = row.get("Vietnamese")
            if orig is None or orig == "nan":
                orig = row.get("Original")
            clean = row.get("Vietnamese_Cleaned")
            if clean is None or clean == "nan":
                clean = row.get("Cleaned")

            orig = str(orig).strip() if orig is not None else ""
            clean = str(clean).strip() if clean is not None else ""
            if not orig or orig == "nan": continue
            result[orig] = clean if clean != "nan" else ""

        with open(CLEANV_JSON, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=4, ensure_ascii=False)
        return True, "Success"
    except Exception as e:
        return False, str(e)

def sync_para_template():
    """Syncs ParaTemplate.xlsx to para_template.json"""
    if not os.path.exists(PARA_TEMPLATE_XLSX):
        return False, f"Missing {PARA_TEMPLATE_XLSX}"
    try:
        df = pd.read_excel(PARA_TEMPLATE_XLSX).astype(str)
        result = {}
        for _, row in df.iterrows():
            vn = row.get("Vietnamese", "").strip()
            if not vn or vn == "nan": continue
            result[vn] = {
                "E": row.get("E", "").strip() if row.get("E") != "nan" else "",
                "Hs": row.get("Hs", "").strip() if row.get("Hs") != "nan" else "",
                "Ht": row.get("Ht", "").strip() if row.get("Ht") != "nan" else ""
            }
        with open(PARA_TEMPLATE_JSON, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=4, ensure_ascii=False)
        return True, "Success"
    except Exception as e:
        return False, str(e)

def sync_dictionary_v3():
    """Syncs Dictionary_v3.xlsx to dictionary_v3.json"""
    if not os.path.exists(DICTIONARY_V3_XLSX):
        return False, f"Missing {DICTIONARY_V3_XLSX}"
    try:
        df = pd.read_excel(DICTIONARY_V3_XLSX).astype(str)
        result = df.to_dict(orient="records")
        with open(DICTIONARY_V3_JSON, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=4, ensure_ascii=False)
        return True, "Success"
    except Exception as e:
        return False, str(e)

def sync_all_templates():
    """Syncs all three templates and returns results summary."""
    results = {}
    results["CleanV"] = sync_clean_v()
    results["ParaTemplate"] = sync_para_template()
    results["DictionaryV3"] = sync_dictionary_v3()
    return results


def load_and_fill_v3_dictionary(metadata):
    """
    Loads dictionary_v3.json and replaces all metadata tags.
    If JSON is missing, syncs from XLSX first.
    Returns a resolved DataFrame.
    """
    # 0. Ensure JSON exists
    if not os.path.exists(DICTIONARY_V3_JSON):
        if os.path.exists(DICTIONARY_V3_FILE):
            sync_dictionary_v3()
        else:
            return None
            
    try:
        # 1. Load from JSON
        with open(DICTIONARY_V3_JSON, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        # Convert to DataFrame (temporarily) to use the existing tag-filling logic
        # This keeps the logic robust and minimizes changes to the core replacement loops
        df = pd.DataFrame(data)
        
        # 2. Build substitution map (normalized to NFC)
        sub_map = get_metadata_substitution_map(metadata)
        
        # 3. Force every column to string for total reliability
        df = df.astype(str)
        
        # 4. Explicit Column-by-Column, Tag-by-Tag Replacement
        # This approach is chosen for maximum compatibility and robustness
        
        # Tags that should NOT be replaced in the Vietnamese column (because they are for ParaTemplate)
        protected_tag_names = [
            "têncôngty", "tên công ty",
            "ngàykếtthúcnăm", "ngày kết thúc năm",
            "ngàybáocáo", "ngày báo cáo",
            "nămbáocáo", "năm báo cáo"
        ]
        # Normalize to NFC for consistent matching check
        protected_tag_names = [unicodedata.normalize('NFC', t.lower()) for t in protected_tag_names]

        # English month names for Column E (Column B)
        english_months = {
            1: "January", 2: "February", 3: "March", 4: "April",
            5: "May", 6: "June", 7: "July", 8: "August",
            9: "September", 10: "October", 11: "November", 12: "December"
        }

        for tag_with_brackets, val in sub_map.items():
            if not tag_with_brackets: continue
            
            # Extract tag name: [V_NAME] -> V_NAME
            # We must preserve casing for the regex to match [V_NAME] specifically
            tag_name = tag_with_brackets.replace("[", "").replace("]", "").strip()
            tag_name_nfc = unicodedata.normalize('NFC', tag_name)
            
            # Check if this tag is protected in Vietnamese column (case-insensitive check)
            is_protected = tag_name_nfc.lower() in protected_tag_names
            
            # Build robust regex: matches [ whitespace tag_name whitespace ]
            # CASE-SENSITIVE matching to support v_name vs V_NAME
            pattern = re.compile(r"\[\s*" + re.escape(tag_name_nfc) + r"\s*\]")
            
            for col in df.columns:
                # Skip 'Vietnamese' column if the tag is protected for ParaTemplate
                if is_protected and str(col).lower() == 'vietnamese':
                    continue
                
                replacement_val = str(val) if val is not None else ""
                
                # SPECIAL: Use English month names if column is 'E' and tag is a month tag
                if str(col).upper() == 'E' and "_month]" in tag_with_brackets.lower():
                    try:
                        m_num = int(replacement_val)
                        if m_num in english_months:
                            replacement_val = english_months[m_num]
                    except: pass
                    
                # Use Series.str.replace with regex=True
                df[col] = df[col].str.replace(pattern, replacement_val, regex=True)
                
        # 4. Final cleaning for every cell (collapsing whitespace, NFC normalization)
        for col in df.columns:
            df[col] = df[col].apply(lambda x: clean_text(x))
            
        return df
    except Exception as e:
        print(f"Error processing Dictionary_v3: {e}")
        return None


def prepare_translation_list(translation_map, case_threshold=25):
    """
    Pre-processes and sorts translation terms for efficiency.
    """
    # Sort keys by length descending to match longest phrases first
    sorted_keys = sorted(translation_map.keys(), key=len, reverse=True)
    prepared = []
    for key in sorted_keys:
        val = translation_map[key]
        if len(key) >= case_threshold:
            # Pre-compile regex for long phrases
            pattern = re.compile(re.escape(key), re.IGNORECASE)
            prepared.append((True, pattern, val))
        else:
            # Keep as string for fast replacement
            prepared.append((False, key, val))
    return prepared

def apply_translations_to_text(text, prepared_list):
    """
    Core string replacement logic using the pre-compiled translation list.
    """
    if not text:
        return text, False
        
    new_text = text
    changed = False
    
    for is_regex, key_or_pattern, val in prepared_list:
        if is_regex:
            # Case-insensitive for long phrases (pre-compiled regex)
            if key_or_pattern.search(new_text):
                new_text = key_or_pattern.sub(val, new_text)
                changed = True
        else:
            # Case-sensitive for short terms (literal replacement)
            if key_or_pattern in new_text:
                new_text = new_text.replace(key_or_pattern, val)
                changed = True
    return new_text, changed

def is_meaningful_text(text):
    """
    Checks if a string contains alphanumeric characters or Vietnamese text.
    Used to distinguish actual content runs from 'marker' runs (dots, symbols).
    """
    if not text:
        return False
    # Check if contains letters or numbers
    if any(c.isalnum() for c in text):
        return True
    # Check if contains Vietnamese characters
    if contains_vietnamese(text):
        return True
    return False

def apply_translations_to_paragraph(paragraph, prepared_list, preserve_newlines=False):
    """
    Applies a pre-processed list of translations to a paragraph.
    Works for single-paragraph units.
    """
    inline = paragraph.runs
    if not inline and not paragraph.text:
        return False

    full_text = "".join(run.text for run in inline)
    
    # Normalize document text to ensure matching with dictionary
    # Pass through the newline preservation flag
    full_text = clean_text(full_text, preserve_newlines=preserve_newlines)
    
    new_text, changed = apply_translations_to_text(full_text, prepared_list)
    
    if changed:
        # Update runs while preserving paragraph-level formatting
        # We try to find the 'best' run to preserve its color/boldness/font
        # 1st Priority: Run with meaningful text (letters/numbers)
        # 2rd Priority: First non-empty run
        format_source_run = None
        
        # Priority 1: Meaningful text
        for run in inline:
            if is_meaningful_text(run.text):
                format_source_run = run
                break
        
        # Priority 2: Any non-empty text
        if not format_source_run:
            for run in inline:
                if run.text and run.text.strip():
                    format_source_run = run
                    break
        
        # Fallback to first run
        if not format_source_run and len(inline) > 0:
            format_source_run = inline[0]
                
        if format_source_run:
            for i, run in enumerate(inline):
                if run == format_source_run:
                    run.text = new_text
                else:
                    run.text = ""
        else:
            paragraph.add_run(new_text)
    return changed

def replace_text_in_paragraph(paragraph, translation_map, case_threshold=25, preserve_newlines=False):
    """
    Backwards compatibility wrapper for apply_translations_to_paragraph.
    """
    prepared = prepare_translation_list(translation_map, case_threshold)
    return apply_translations_to_paragraph(paragraph, prepared, preserve_newlines=preserve_newlines)

def _process_container(container, prepared_list, replaced_paragraphs=None):
    """
    Enhanced processing of containers (Document, Header, Footer).
    Matches multi-paragraph terms inside table cells and text boxes by 
    joining them into a single logical unit.
    """
    count = 0
    element = container._element
    processed_p_elements = set()
    
    if replaced_paragraphs is None:
        replaced_paragraphs = set()

    # 1. Process Multi-paragraph Containers (Cells, Textboxes)
    # tc: Table Cell, txbxContent: Textbox Content
    containers = element.xpath('.//*[local-name()="tc"] | .//*[local-name()="txbxContent"]')
    
    for c_el in containers:
        # Find paragraphs directly within this container (or descendants)
        p_elements = c_el.xpath('.//*[local-name()="p"]')
        if not p_elements: continue
        
        # Build full text by joining paragraph contents
        p_texts = []
        for p_el in p_elements:
            t_nodes = p_el.xpath('.//*[local-name()="t"]')
            p_texts.append("".join(t.text for t in t_nodes if t.text))
        
        full_text = " ".join(p_texts)
        cleaned_body_text = clean_text(full_text)
        
        # --- Signature Block Protection ---
        # If the cell looks like a signature block, we process paragraphs individually
        # to preserve line breaks and formatting.
        is_signature = False
        lower_text = cleaned_body_text.lower()
        
        # Condition A: Representing + Company Name (VN / EN / CN)
        if ("thay mặt và đại diện cho" in lower_text or "on behalf of" in lower_text or "代表" in lower_text) and \
           ("công ty tnhh kiểm toán u&i" in lower_text or "u&i auditing" in lower_text or "u&i 审计" in lower_text):
            is_signature = True
        # Condition B: Auditor + License ID (VN / EN / CN)
        elif ("kiểm toán viên" in lower_text or "auditor" in lower_text or "注册会计师" in lower_text) and \
             ("số giấy cn đkhn kiểm toán" in lower_text or "practising certificate no" in lower_text or "证书编号" in lower_text):
            is_signature = True
        # Condition C: Generic Director / Manager Signature Block
        elif "giám đốc" in lower_text or "tổng giám đốc" in lower_text or "director" in lower_text or "总经理" in lower_text:
            if len(p_elements) <= 10: # Only if it's a reasonably small block (like a signature)
                is_signature = True
            
        if is_signature:
            # Process paragraphs individually to keep the signature layout
            for p_el in p_elements:
                p = Paragraph(p_el, container)
                # ENABLE newline preservation for signature block paragraphs
                if apply_translations_to_paragraph(p, prepared_list, preserve_newlines=True):
                    count += 1
        else:
            # Standard Consolidation Logic (for Split headers like "Mã\nsố")
            new_text, changed = apply_translations_to_text(cleaned_body_text, prepared_list)
            
            if changed:
                # Found a match across the entire container unit (e.g. "Mã\nsố")
                # We consolidate the translation into a single paragraph and clear others.
                # To preserve formatting, we find the 'best' paragraph/run to host the result.
                
                target_para = None
                # Priority 1: Find paragraph with meaningful text
                for p_el in p_elements:
                    p_obj = Paragraph(p_el, container)
                    if is_meaningful_text(p_obj.text):
                        target_para = p_obj
                        break
                
                # Priority 2: Use first paragraph
                if not target_para:
                    target_para = Paragraph(p_elements[0], container)
                
                # Apply translation to the chosen paragraph using run-aware logic
                # We use a case_threshold=9999 (ignore delta) to ENSURE literal match.
                # FIX: We must match AGAINST the target_para's own text, not the combined text,
                # to ensure apply_translations_to_paragraph actually finds something to replace.
                current_p_text = clean_text(target_para.text)
                if current_p_text:
                    dummy_map = prepare_translation_list({current_p_text: new_text}, case_threshold=9999)
                    apply_translations_to_paragraph(target_para, dummy_map)
                else:
                    # Fallback for empty paragraphs: just add a run
                    target_para.add_run(new_text)
                    
                # Remove all other paragraphs in this container to prevent duplicates/layout issues
                for other_p_el in p_elements:
                    if other_p_el != target_para._element:
                        parent = other_p_el.getparent()
                        if parent is not None:
                            parent.remove(other_p_el)
                count += 1
            
        # Mark these paragraphs as handled
        for p_el in p_elements:
            processed_p_elements.add(p_el)
            
    # 2. Process lone paragraphs (not inside cells or textboxes)
    all_p_elements = element.xpath('.//*[local-name()="p"]')
    for p_el in all_p_elements:
        if p_el not in processed_p_elements:
            p = Paragraph(p_el, container)
            
            # Process dictionary translation for all paragraphs
            # (Previously we skipped replaced paragraphs, which blocked placeholders)
            if apply_translations_to_paragraph(p, prepared_list):
                count += 1
                
    return count





def set_document_default_fonts(doc, target_col):
    """
    Sets the document-level default fonts and size (DocDefaults).
    Equivalent to clicking 'Set As Default' in Word.
    Specifically sets DFKai-SB for Chinese and Times New Roman for Latin.
    """
    if target_col not in ["Hs", "Ht"]:
        return

    # Access styles part
    styles_element = doc.styles.element
    
    # 1. Ensure docDefaults exists
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, doc_defaults)
        
    # 2. Ensure rPrDefault exists
    r_pr_default = doc_defaults.find(qn('w:rPrDefault'))
    if r_pr_default is None:
        r_pr_default = OxmlElement('w:rPrDefault')
        doc_defaults.append(r_pr_default)
        
    # 3. Ensure rPr exists
    r_pr = r_pr_default.find(qn('w:rPr'))
    if r_pr is None:
        r_pr = OxmlElement('w:rPr')
        r_pr_default.append(r_pr)
        
    # 4. Set Fonts
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    
    r_fonts.set(qn('w:ascii'), "Times New Roman")
    r_fonts.set(qn('w:hAnsi'), "Times New Roman")
    r_fonts.set(qn('w:eastAsia'), "DFKai-SB")
    r_fonts.set(qn('w:cs'), "Times New Roman")
    
    # 5. Set Size (10pt = 20 half-points)
    sz = r_pr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        r_pr.append(sz)
    sz.set(qn('w:val'), '20')
    
    sz_cs = r_pr.find(qn('w:szCs'))
    if sz_cs is None:
        sz_cs = OxmlElement('w:szCs')
        r_pr.append(sz_cs)
    sz_cs.set(qn('w:val'), '20')
    
    # 6. Set Language
    lang = r_pr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        r_pr.append(lang)
    lang_val = "zh-CN" if target_col == "Hs" else "zh-TW"
    lang.set(qn('w:eastAsia'), lang_val)

def _get_all_paragraphs(doc):
    """
    Helper to yield all paragraphs in the document (body, tables, headers, footers).
    """
    # 1. Main body paragraphs
    for para in doc.paragraphs:
        yield para
        
    # 2. Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
                    
    # 3. Headers and Footers
    for section in doc.sections:
        containers = [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer
        ]
        try:
            containers.extend([section.even_page_header, section.even_page_footer])
        except: pass
        
        for container in containers:
            if container:
                for para in container.paragraphs:
                    yield para
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                yield para

def replace_text_in_document(doc, translation_map, case_threshold=25, cleanv_map=None, para_map=None, target_col="E", metadata=None, process_settings=None):
    """
    Performs global search and replace in paragraphs, tables, headers and footers.
    Respects process_settings toggles.
    """
    # Default to "All ON" if no settings provided
    if process_settings is None:
        process_settings = {k: True for k in ["metadata", "unicode", "clean_v", "para_template", "dictionary", "dual_font", "table_size", "date_format", "textbox", "highlight", "suggestion"]}

    # 0. Load maps if not provided but exist
    if cleanv_map is None:
        cleanv_map = load_cleanv_map()
    if para_map is None:
        para_map = load_para_template_map()
        
    # NEW: Capture Snapshot of Original Paragraph Texts for downstream fuzzy suggestions
    original_texts = {p._element: p.text for p in _get_all_paragraphs(doc)}

    total_count = 0

    # Step 1: Explicit Unicode Normalization (Always first for consistent matching)
    if process_settings.get("unicode", True):
        apply_unicode_normalization(doc)

    # Step 4: Paragraph Template Replacements (Full paragraph swaps)
    # Must run BEFORE metadata and CleanV so that original phrases can match correctly.
    replaced_paras = set()
    if process_settings.get("para_template", True) and para_map:
        _, replaced_paras = apply_paragraph_templates(doc, para_map, target_col)

    # Step 2: Global Normalization (CleanV Typo Correction)
    if process_settings.get("clean_v", True) and cleanv_map:
        apply_cleanv_normalization(doc, cleanv_map)



    # Step 0.5: Set Document Defaults (Set As Default)
    # Applied to Chinese reports to force DFKai-SB for any unstyled text.
    if target_col in ["Hs", "Ht"]:
        set_document_default_fonts(doc, target_col)

    # Pass 3: Dictionary-based replacements (Final translation)
    if process_settings.get("dictionary", True):
        prepared_list = prepare_translation_list(translation_map, case_threshold)
        
        # Add normalization list (CleanV) to the start of prepared_list ONLY if clean_v IS enabled
        if process_settings.get("clean_v", True) and cleanv_map:
            norm_list = []
            for key in cleanv_map:
                norm_list.append((False, key, cleanv_map[key]))
            prepared_list = norm_list + prepared_list

        # Process the main document body
        total_count += _process_container(doc, prepared_list, replaced_paragraphs=replaced_paras)
        
        # Process all headers and footers in all sections
        for section in doc.sections:
            total_count += _process_container(section.header, prepared_list, replaced_paragraphs=replaced_paras)
            total_count += _process_container(section.footer, prepared_list, replaced_paragraphs=replaced_paras)
            if section.different_first_page_header_footer:
                total_count += _process_container(section.first_page_header, prepared_list, replaced_paragraphs=replaced_paras)
                total_count += _process_container(section.first_page_footer, prepared_list, replaced_paragraphs=replaced_paras)
            try:
                total_count += _process_container(section.even_page_header, prepared_list, replaced_paragraphs=replaced_paras)
                total_count += _process_container(section.even_page_footer, prepared_list, replaced_paragraphs=replaced_paras)
            except: pass

    # Step 4: Final Chinese Currency Cleanup (Hs/Ht only)
    if process_settings.get("dictionary", True) and target_col in ["Hs", "Ht"]:
        apply_chinese_currency_cleanup(doc)
        
    # Step 5: Format dates in tables
    if process_settings.get("date_format", True):
        format_dates_in_tables(doc, target_col)
    
    # Step 7: Dual-font formatting (Chinese only)
    if process_settings.get("dual_font", True) and target_col in ["Hs", "Ht"]:
        apply_chinese_font_formatting(doc, target_col)
        
    # Step 7.5: Financial Number Separator Swap (. to , and , to .)
    # Only for non-Vietnamese reports. Safe for Link Fields.
    if process_settings.get("number_swap", True) and target_col != "V":
        apply_financial_number_formatting(doc, target_col)

    # Step 8: Specialized Table Sizing and Layout
    if process_settings.get("table_size", True):
        apply_sizing_and_layout(doc, target_col)
    
    # Step 9: Specialized TextBox/Draft Handling
    if process_settings.get("textbox", True):
        apply_special_textbox_formatting(doc, target_col)
    
    # Step 12: Signer Accent Removal
    # New logic: If signer name has accents, replace with unaccented in whole doc.
    if process_settings.get("signer_accents", True) and metadata:
        apply_signer_accent_removal(doc, metadata)

    # Step 10 & 11: Highlight and Suggest
    if process_settings.get("highlight", True) or process_settings.get("suggestion", True):
        highlight_vietnamese_text(
            doc, 
            translation_map, 
            original_texts=original_texts,
            show_suggestions=process_settings.get("suggestion", True)
        )

    return total_count
